"""
Clean SEO Content Automation Service - 7 Step Process
Integrates with Django for streamlined blog generation
"""
import os
import sys
import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Tuple

# Add the parent directory to sys.path to import the original script
parent_dir = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(parent_dir))

from django.conf import settings
from django.core.files.base import ContentFile
import tempfile

logger = logging.getLogger(__name__)


def parse_blog_titles(titles_input: str, max_titles: int = 5) -> Tuple[List[str], Dict[str, Any]]:
    """Parse comma-separated blog titles into a list"""
    parsing_info = {
        'original_count': 0,
        'processed_count': 0,
        'duplicates_removed': 0,
        'empty_removed': 0,
        'truncated': False,
        'warnings': []
    }
    
    if not titles_input or not titles_input.strip():
        return [], parsing_info
    
    # Split by commas and clean
    raw_titles = [title.strip() for title in titles_input.split(',')]
    parsing_info['original_count'] = len(raw_titles)
    
    # Filter out empty titles
    non_empty_titles = [title for title in raw_titles if title]
    parsing_info['empty_removed'] = len(raw_titles) - len(non_empty_titles)
    
    # Remove duplicates while preserving order
    seen = set()
    unique_titles = []
    for title in non_empty_titles:
        title_lower = title.lower()
        if title_lower not in seen:
            seen.add(title_lower)
            unique_titles.append(title)
        else:
            parsing_info['duplicates_removed'] += 1
    
    # Limit the number of titles
    if len(unique_titles) > max_titles:
        parsing_info['truncated'] = True
        unique_titles = unique_titles[:max_titles]
    
    parsing_info['processed_count'] = len(unique_titles)
    
    # Validate title length
    for i, title in enumerate(unique_titles):
        if len(title) > 200:
            unique_titles[i] = title[:200] + "..."
    
    return unique_titles, parsing_info


class SEOBlogGenerator:
    """Clean SEO blog generator following the 7-step process"""
    
    def __init__(self):
        self.media_root = getattr(settings, 'MEDIA_ROOT', tempfile.gettempdir())
        # Initialize Gemini model once
        self.llm = None
        if os.getenv('GOOGLE_GEMINI_API_KEY'):
            from langchain_google_genai import ChatGoogleGenerativeAI
            self.llm = ChatGoogleGenerativeAI(
                model="gemini-2.5-flash",
                google_api_key=os.getenv("GOOGLE_GEMINI_API_KEY"),
                temperature=0.3,
                max_tokens=8192
            )

    def generate_blog(self, title, primary_keywords=None, num_competitors=3, secondary_keywords=None, blog_outline=None, target_length=None):
        """
        Clean 7-step SEO blog generation process:
        1. Get top 3 competitors and PAA questions
        2. Scrape competitor content  
        3. Analyze content gaps
        4. Extract competitor keywords
        5. Analyze content structure
        6. Generate detailed blog outline
        7. Generate final SEO optimized blog
        """
        try:
            # Check API keys
            if not os.getenv('GOOGLE_GEMINI_API_KEY'):
                return {'success': False, 'error': 'GOOGLE_GEMINI_API_KEY is required', 'content': None}
            if not os.getenv('SERP_API_KEY'):
                return {'success': False, 'error': 'SERP_API_KEY is required', 'content': None}
            
            # Parse target word count
            target_word_count = None
            if target_length:
                import re
                numeric_values = re.findall(r'\d+', target_length)
                if numeric_values:
                    target_word_count = int(numeric_values[0])
            
            logger.info(f"Starting 7-step blog generation for: {title}")
            
            # STEP 1: Get competitors and PAA questions
            competitors, paa_questions = self._get_competitors_and_paa(title, num_competitors)
            if not competitors:
                return {'success': False, 'error': 'No competitors found', 'content': None}
            
            # STEP 2: Scrape competitor content
            scraped_data = self._scrape_competitors(competitors)
            if not scraped_data:
                return {'success': False, 'error': 'Failed to scrape competitor content', 'content': None}
            
            # STEP 3: Analyze content gaps
            content_gaps = self._analyze_content_gaps(scraped_data, paa_questions, title)
            
            # STEP 4: Extract competitor keywords
            competitor_keywords = self._extract_competitor_keywords(scraped_data, title)
            
            # STEP 5: Analyze content structure
            content_structure = self._analyze_content_structure(scraped_data)
            
            # STEP 6: Generate blog outline
            blog_outline = self._generate_blog_outline(title, content_gaps, competitor_keywords, content_structure, target_word_count)
            
            # STEP 7: Generate final blog
            final_blog = self._generate_final_blog(title, blog_outline, target_word_count)
            
            # Normalize output to plain markdown
            formatted_post = self._normalize_model_output_to_markdown(final_blog)
            
            return {
                'success': True,
                'error': None,
                'content': {
                    'formatted_post': formatted_post,
                    'blog_post_data': {'blog_post': {'content': formatted_post, 'title': title}},
                    'content_strategy': {'content_gaps': content_gaps, 'keywords': competitor_keywords},
                    'word_count': f"{len(formatted_post.split())} words",
                    'target_keyword': title,
                    'competitors_analyzed': len(scraped_data)
                }
            }
            
        except Exception as e:
            logger.error(f"Blog generation failed: {str(e)}")
            return {'success': False, 'error': f"Blog generation failed: {str(e)}", 'content': None}

    # ==================== 7-STEP PROCESS METHODS ====================
    
    def _get_competitors_and_paa(self, title, num_competitors=3):
        """STEP 1: Get top competitors and PAA questions using SERP API"""
        try:
            from seo_content_automation import get_top_5_links_and_paa
            competitors, paa_questions = get_top_5_links_and_paa(title, os.getenv("SERP_API_KEY"), num_competitors)
            logger.info(f"Found {len(competitors)} competitors and {len(paa_questions)} PAA questions")
            return competitors, paa_questions
        except Exception as e:
            logger.error(f"Error getting competitors: {e}")
            return [], []
    
    def _scrape_competitors(self, competitors):
        """STEP 2: Scrape content from competitors"""
        try:
            from seo_content_automation import scrape_blog
            scraped_data = []
            for i, url in enumerate(competitors, 1):
                logger.info(f"Scraping competitor {i}: {url}")
                data = scrape_blog(url)
                if data:
                    scraped_data.append(data)
            logger.info(f"Successfully scraped {len(scraped_data)} competitors")
            return scraped_data
        except Exception as e:
            logger.error(f"Error scraping competitors: {e}")
            return []
    
    def _analyze_content_gaps(self, scraped_data, paa_questions, title):
        """STEP 3: Analyze content gaps using Gemini"""
        if not self.llm:
            return {"error": "Gemini not available"}
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            # Prepare competitor summary
            competitor_summary = ""
            for i, blog in enumerate(scraped_data, 1):
                competitor_summary += f"\n--- Competitor {i} ---\n"
                competitor_summary += f"Title: {blog.get('title', 'N/A')}\n"
                competitor_summary += f"H2 Headings: {[h['heading'] for h in blog.get('h2', [])]}\n"
            
            paa_summary = "\n".join([f"- {q}" for q in paa_questions])
            
            system_prompt = """Analyze competitor content and identify content gaps. Return JSON:
            {
                "content_gaps": ["gap1", "gap2", "gap3"],
                "missing_topics": ["topic1", "topic2"],
                "opportunities": ["opportunity1", "opportunity2"]
            }"""
            
            human_prompt = f"""Target: "{title}"
            
            Competitors: {competitor_summary}
            
            PAA Questions: {paa_summary}
            
            Identify content gaps and opportunities to create better content."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            try:
                import json
                return json.loads(response.content)
            except:
                return {"analysis": response.content}
                
        except Exception as e:
            logger.error(f"Error analyzing content gaps: {e}")
            return {"error": str(e)}
    
    def _extract_competitor_keywords(self, scraped_data, title):
        """STEP 4: Extract keywords from competitors using Gemini"""
        if not self.llm:
            return {"error": "Gemini not available"}
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            # Prepare content for analysis
            content_text = ""
            for blog in scraped_data:
                content_text += f"Title: {blog.get('title', '')}\n"
                content_text += f"Meta: {blog.get('metadescription', '')}\n"
                for h in blog.get('h2', []):
                    content_text += f"H2: {h['heading']}\n"
            
            system_prompt = """Extract keywords from competitor content. Return JSON:
            {
                "primary_keywords": ["keyword1", "keyword2"],
                "long_tail_keywords": ["long tail 1", "long tail 2"],
                "secondary_keywords": ["secondary1", "secondary2"]
            }"""
            
            human_prompt = f"""Target: "{title}"
            
            Competitor Content: {content_text[:2000]}
            
            Extract relevant keywords for SEO optimization."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            try:
                import json
                return json.loads(response.content)
            except:
                return {"keywords": response.content}
                
        except Exception as e:
            logger.error(f"Error extracting keywords: {e}")
            return {"error": str(e)}
    
    def _analyze_content_structure(self, scraped_data):
        """STEP 5: Analyze competitor content structure using Gemini"""
        if not self.llm:
            return {"error": "Gemini not available"}
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            # Analyze structure
            structure_info = ""
            for i, blog in enumerate(scraped_data, 1):
                structure_info += f"\nCompetitor {i}:\n"
                structure_info += f"H1 count: {len(blog.get('h1', []))}\n"
                structure_info += f"H2 count: {len(blog.get('h2', []))}\n"
                structure_info += f"H3 count: {len(blog.get('h3', []))}\n"
            
            system_prompt = """Analyze content structure and recommend optimal structure. Return JSON:
            {
                "recommended_structure": {
                    "h1_count": 1,
                    "h2_count": 5,
                    "h3_count": 8,
                    "word_count_range": "2000-3000"
                },
                "content_patterns": ["pattern1", "pattern2"]
            }"""
            
            human_prompt = f"""Competitor Structure Analysis: {structure_info}
            
            Recommend optimal content structure for better SEO."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            try:
                import json
                return json.loads(response.content)
            except:
                return {"structure": response.content}
                
        except Exception as e:
            logger.error(f"Error analyzing structure: {e}")
            return {"error": str(e)}
    
    def _generate_blog_outline(self, title, content_gaps, keywords, structure, target_word_count):
        """STEP 6: Generate detailed blog outline using Gemini"""
        if not self.llm:
            return {"error": "Gemini not available"}
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            word_count_text = f"Target: {target_word_count} words" if target_word_count else "Target: 2000-3000 words"
            
            system_prompt = f"""Create a detailed blog outline. Return JSON:
            {{
                "title": "{title}",
                "sections": [
                    {{
                        "heading": "Section heading",
                        "word_count": "300-400",
                        "key_points": ["point1", "point2"]
                    }}
                ],
                "estimated_word_count": "{word_count_text}"
            }}"""
            
            human_prompt = f"""Title: "{title}"
            
            Content Gaps: {content_gaps}
            Keywords: {keywords}
            Structure Analysis: {structure}
            {word_count_text}
            
            Create a comprehensive blog outline that addresses content gaps and uses relevant keywords."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            try:
                import json
                return json.loads(response.content)
            except:
                return {"outline": response.content}
                
        except Exception as e:
            logger.error(f"Error generating outline: {e}")
            return {"error": str(e)}
    
    def _generate_final_blog(self, title, outline, target_word_count):
        """STEP 7: Generate final SEO optimized blog using Gemini"""
        if not self.llm:
            return f"# {title}\n\nBlog generation failed - Gemini not available."
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            word_count_instruction = f"Write exactly {target_word_count} words." if target_word_count else "Write 2000-3000 words."
            
            system_prompt = f"""You are an expert SEO content writer. Write a complete, high-quality blog post in markdown format.
            
            Requirements:
            - {word_count_instruction}
            - SEO optimized with natural keyword usage
            - Engaging and informative content
            - Proper heading structure (H1, H2, H3)
            - Include introduction and conclusion
            - Actionable insights and value for readers
            
            Return ONLY the blog content in markdown format, no JSON or extra text."""
            
            human_prompt = f"""Write a complete blog post with title: "{title}"
            
            Blog Outline: {outline}
            
            Create comprehensive, SEO-friendly content that follows the outline and provides real value to readers."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            return response.content
            
        except Exception as e:
            logger.error(f"Error generating final blog: {e}")
            return f"# {title}\n\nBlog generation failed: {str(e)}"

    def _normalize_model_output_to_markdown(self, raw_text: str) -> str:
        """Normalize model output to clean markdown text"""
        try:
            if not raw_text:
                return ''

            text = raw_text.strip()

            # Remove code fences
            if text.startswith('```') and text.endswith('```'):
                parts = text.split('\n')
                if parts and parts[0].startswith('```'):
                    parts = parts[1:]
                if parts and parts[-1].startswith('```'):
                    parts = parts[:-1]
                text = '\n'.join(parts).strip()

            # Try to extract from JSON if present
            if text.startswith('{'):
                try:
                    import json
                    parsed = json.loads(text)
                    if isinstance(parsed, dict):
                        if 'final_article' in parsed:
                            return parsed['final_article'].strip()
                        elif 'blog_post' in parsed and 'content' in parsed['blog_post']:
                            return parsed['blog_post']['content'].strip()
                except:
                    pass

            return text
        except:
            return raw_text if isinstance(raw_text, str) else str(raw_text)

    def save_as_document(self, formatted_post, filename):
        """Save the formatted blog post as a .docx file"""
        try:
            from docx import Document
            import re
            
            file_path = os.path.join(self.media_root, 'generated_blogs', filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            doc = Document()
            
            lines = formatted_post.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                else:
                    # Clean markdown formatting
                    clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)  # Remove bold
                    clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)  # Remove italic
                    clean_text = re.sub(r'`([^`]+)`', r'\1', clean_text)  # Remove code
                    doc.add_paragraph(clean_text)
            
            doc.save(file_path)
            logger.info(f"Document saved to: {file_path}")
            return file_path
                
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return None
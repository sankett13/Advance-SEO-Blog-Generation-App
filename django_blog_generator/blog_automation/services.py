"""
Clean SEO Content Automation Service - 7 Step Process
Integrates with Django for streamlined blog generation
"""
import os
import sys
import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Tuple
import dotenv

# Load environment variables from .env file
dotenv.load_dotenv()

# Add the parent directory to sys.path to import the original script
parent_dir = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(parent_dir))

from django.conf import settings
from django.core.files.base import ContentFile
import tempfile

logger = logging.getLogger(__name__)


def check_api_configuration():
    """Check if required API keys are configured"""
    gemini_key = os.getenv('GOOGLE_GEMINI_API_KEY')
    serp_key = os.getenv('SERP_API_KEY')
    
    all_configured = bool(gemini_key and serp_key)
    missing_keys = []
    
    if not gemini_key:
        missing_keys.append('GOOGLE_GEMINI_API_KEY')
    if not serp_key:
        missing_keys.append('SERP_API_KEY')
    
    status = {
        'gemini_configured': bool(gemini_key),
        'serp_configured': bool(serp_key),
        'all_configured': all_configured,
        'configured': all_configured,  # Template compatibility
        'missing_keys': missing_keys
    }
    
    return status


def parse_blog_titles(titles_input: str, max_titles: int = 5) -> Tuple[List[str], Dict[str, Any]]:
    parsing_info = {
        'original_count': 0,
        'processed_count': 0,
        'duplicates_removed': 0,
        'empty_removed': 0,
        'truncated': False,
        'warnings': []
    }
    
    if not titles_input or not titles_input.strip():
        return [], parsing_info
    
    # Split by commas and clean
    raw_titles = [title.strip() for title in titles_input.split(',')]
    parsing_info['original_count'] = len(raw_titles)
    
    # Filter out empty titles
    non_empty_titles = [title for title in raw_titles if title]
    parsing_info['empty_removed'] = len(raw_titles) - len(non_empty_titles)
    
    # Remove duplicates while preserving order
    seen = set()
    unique_titles = []
    for title in non_empty_titles:
        title_lower = title.lower()
        if title_lower not in seen:
            seen.add(title_lower)
            unique_titles.append(title)
        else:
            parsing_info['duplicates_removed'] += 1
    
    # Limit the number of titles
    if len(unique_titles) > max_titles:
        parsing_info['truncated'] = True
        unique_titles = unique_titles[:max_titles]
    
    parsing_info['processed_count'] = len(unique_titles)
    
    # Validate title length
    for i, title in enumerate(unique_titles):
        if len(title) > 200:
            unique_titles[i] = title[:200] + "..."
    
    return unique_titles, parsing_info


class SEOBlogGenerator:
    """Clean SEO blog generator following the 7-step process"""
    
    def __init__(self):
        self.media_root = getattr(settings, 'MEDIA_ROOT', tempfile.gettempdir())
        # Initialize Gemini model once
        self.llm = None
        if os.getenv('GOOGLE_GEMINI_API_KEY'):
            from langchain_google_genai import ChatGoogleGenerativeAI
            self.llm = ChatGoogleGenerativeAI(
                model="gemini-2.5-flash",
                google_api_key=os.getenv("GOOGLE_GEMINI_API_KEY"),
                temperature=0.3,
                max_tokens=8192
            )

    def generate_multiple_blogs(self, titles_input, primary_keywords=None, num_competitors=3, 
                               secondary_keywords=None, blog_outline=None, target_length=None, max_titles=5):
        """
        Generate multiple blogs from comma-separated titles (max 5 at a time)
        
        Args:
            titles_input (str): Comma-separated blog titles
            max_titles (int): Maximum number of titles to process (default: 5)
            Other args: Same as generate_blog method
            
        Returns:
            dict: Batch processing results with individual blog results
        """
        # Parse the titles input
        blog_titles, parsing_info = parse_blog_titles(titles_input, max_titles)
        
        if not blog_titles:
            return {
                'success': False,
                'error': 'No valid blog titles provided',
                'blog_results': [],
                'successful_blogs': 0,
                'failed_blogs': 0,
                'parsing_info': parsing_info,
                'errors': ['No valid blog titles found']
            }
        
        logger.info(f"Starting batch generation for {len(blog_titles)} blog titles")
        
        # Initialize results tracking
        blog_results = []
        successful_blogs = 0
        failed_blogs = 0
        batch_errors = []
        
        # Generate each blog
        for i, title in enumerate(blog_titles, 1):
            logger.info(f"Generating blog {i}/{len(blog_titles)}: {title}")
            
            try:
                # Generate individual blog using the same 7-step process
                result = self.generate_blog(
                    title=title.strip(),
                    primary_keywords=primary_keywords,
                    num_competitors=num_competitors,
                    secondary_keywords=secondary_keywords,
                    blog_outline=blog_outline,
                    target_length=target_length
                )
                
                # Add title and batch processing info to result for tracking
                result['title'] = title.strip()
                result['batch_position'] = i
                result['batch_total'] = len(blog_titles)
                blog_results.append(result)
                
                if result['success']:
                    successful_blogs += 1
                    logger.info(f"✅ Blog {i} generated successfully: {title}")
                else:
                    failed_blogs += 1
                    error_msg = f"Blog {i} '{title}': {result.get('error', 'Unknown error')}"
                    batch_errors.append(error_msg)
                    logger.error(f"❌ {error_msg}")
                    
            except Exception as e:
                failed_blogs += 1
                error_msg = f"Blog {i} '{title}': Exception - {str(e)}"
                batch_errors.append(error_msg)
                logger.error(f"❌ {error_msg}")
                
                # Add failed result to maintain consistency
                blog_results.append({
                    'success': False,
                    'error': str(e),
                    'content': None,
                    'title': title.strip()
                })
        
        # Compile batch results
        batch_result = {
            'success': successful_blogs > 0,
            'blog_results': blog_results,
            'successful_blogs': successful_blogs,
            'failed_blogs': failed_blogs,
            'total_processed': len(blog_titles),
            'parsing_info': parsing_info,
            'errors': batch_errors
        }
        
        logger.info(f"Batch generation completed: {successful_blogs} successful, {failed_blogs} failed out of {len(blog_titles)} titles")
        
        return batch_result

    def generate_blog(self, title, primary_keywords=None, num_competitors=3, secondary_keywords=None, blog_outline=None, target_length=None):
        """
        Clean 7-step SEO blog generation process:
        1. Get top 3 competitors and PAA questions
        2. Scrape competitor content  
        3. Analyze content gaps
        4. Extract competitor keywords
        5. Analyze content structure
        6. Generate detailed blog outline
        7. Generate final SEO optimized blog
        """
        try:
            # Check API keys
            if not os.getenv('GOOGLE_GEMINI_API_KEY'):
                return {'success': False, 'error': 'GOOGLE_GEMINI_API_KEY is required', 'content': None}
            if not os.getenv('SERP_API_KEY'):
                return {'success': False, 'error': 'SERP_API_KEY is required', 'content': None}
            
            # Parse target word count
            target_word_count = None
            if target_length:
                import re
                numeric_values = re.findall(r'\d+', target_length)
                if numeric_values:
                    target_word_count = int(numeric_values[0])
            
            logger.info(f"Starting 7-step blog generation for: {title}")
            
            # STEP 1: Get competitors and PAA questions
            print("Competitors and PAA Extraction")
            competitors, paa_questions = self._get_competitors_and_paa(title, num_competitors)
            if not competitors:
                return {'success': False, 'error': 'No competitors found', 'content': None}
            
            # STEP 2: Scrape competitor content
            print("Competitors Scraping")
            scraped_data = self._scrape_competitors(competitors)
            if not scraped_data:
                return {'success': False, 'error': 'Failed to scrape competitor content', 'content': None}
            
            # STEP 3: Analyze content gaps
            print("Content Gap Analysis using Gemini")
            content_gaps = self._analyze_content_gaps(scraped_data, paa_questions, title)
            
            # STEP 4: Extract competitor keywords
            print("Competitor Keyword Extraction using Gemini")
            competitor_keywords = self._extract_competitor_keywords(scraped_data, title)
            
            # STEP 5: Analyze content structure
            print("Competitor Content Structure Analysis using Gemini")
            content_structure = self._analyze_content_structure(scraped_data)
            
            # STEP 6: Generate blog outline
            print("Blog Outline Generation using Gemini")
            blog_outline = self._generate_blog_outline(title, content_gaps, competitor_keywords, content_structure, target_word_count)
            
            # STEP 7: Generate final blog
            print("Final Blog Generation using Gemini")
            final_blog = self._generate_final_blog(title, blog_outline, target_word_count)
            
            # Normalize output to plain markdown
            formatted_post = self._normalize_model_output_to_markdown(final_blog)
            
            # Compile comprehensive competitor analysis data
            comprehensive_analysis = {
                'step_1_competitors_and_paa': {
                    'competitors': competitors,
                    'paa_questions': paa_questions,
                    'num_competitors_found': len(competitors),
                    'num_paa_questions': len(paa_questions)
                },
                'step_2_scraped_data': {
                    'scraped_content': scraped_data,
                    'successful_scrapes': len(scraped_data),
                    'scraping_summary': [
                        {
                            'url': data.get('url', 'N/A'),
                            'title': data.get('title', 'N/A'),
                            'h1_count': len(data.get('h1', [])),
                            'h2_count': len(data.get('h2', [])),
                            'h3_count': len(data.get('h3', [])),
                            'meta_description': data.get('metadescription', 'N/A')[:100] + '...' if data.get('metadescription') else 'N/A'
                        } for data in scraped_data
                    ]
                },
                'step_3_content_gaps': {
                    'analysis': content_gaps,
                    'analysis_type': 'content_gap_identification'
                },
                'step_4_competitor_keywords': {
                    'extracted_keywords': competitor_keywords,
                    'analysis_type': 'keyword_extraction'
                },
                'step_5_content_structure': {
                    'structure_analysis': content_structure,
                    'analysis_type': 'content_structure_optimization'
                },
                'step_6_blog_outline': {
                    'generated_outline': blog_outline,
                    'outline_type': 'seo_optimized_structure',
                    'target_word_count': target_word_count
                },
                'step_7_generation_metadata': {
                    'final_blog_generated': True,
                    'final_word_count': len(formatted_post.split()),
                    'generation_timestamp': datetime.now().isoformat(),
                    'model_used': 'gemini-2.5-flash'
                },
                'process_summary': {
                    'total_steps_completed': 7,
                    'competitors_analyzed': len(scraped_data),
                    'target_keyword': title,
                    'analysis_complete': True,
                    'generation_successful': True
                }
            }
            
            return {
                'success': True,
                'error': None,
                'content': {
                    'formatted_post': formatted_post,
                    'blog_post_data': {'blog_post': {'content': formatted_post, 'title': title}},
                    'content_strategy': {'content_gaps': content_gaps, 'keywords': competitor_keywords},
                    'word_count': f"{len(formatted_post.split())} words",
                    'target_keyword': title,
                    'competitors_analyzed': len(scraped_data),
                    'comprehensive_analysis': comprehensive_analysis
                }
            }
            
        except Exception as e:
            logger.error(f"Blog generation failed: {str(e)}")
            return {'success': False, 'error': f"Blog generation failed: {str(e)}", 'content': None}

    # ==================== 7-STEP PROCESS METHODS ====================
    
    def _get_competitors_and_paa(self, title, num_competitors=3):
        """STEP 1: Get top competitors and PAA questions using SERP API"""
        try:
            from seo_content_automation import get_top_5_links_and_paa
            competitors, paa_questions = get_top_5_links_and_paa(title, os.getenv("SERP_API_KEY"), num_competitors)
            logger.info(f"Found {len(competitors)} competitors and {len(paa_questions)} PAA questions")
            return competitors, paa_questions
        except Exception as e:
            logger.error(f"Error getting competitors: {e}")
            return [], []
    
    def _scrape_competitors(self, competitors):
        """STEP 2: Scrape content from competitors"""
        try:
            from seo_content_automation import scrape_blog
            scraped_data = []
            for i, url in enumerate(competitors, 1):
                logger.info(f"Scraping competitor {i}: {url}")
                data = scrape_blog(url)
                if data:
                    scraped_data.append(data)
            logger.info(f"Successfully scraped {len(scraped_data)} competitors")
            return scraped_data
        except Exception as e:
            logger.error(f"Error scraping competitors: {e}")
            return []
    
    def _analyze_content_gaps(self, scraped_data, paa_questions, title):
        """STEP 3: Analyze content gaps using Gemini"""
        if not self.llm:
            return {"error": "Gemini not available"}
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            # Prepare competitor summary
            competitor_summary = ""
            for i, blog in enumerate(scraped_data, 1):
                competitor_summary += f"\n--- Competitor {i} ---\n"
                competitor_summary += f"Title: {blog.get('title', 'N/A')}\n"
                competitor_summary += f"H2 Headings: {[h['heading'] for h in blog.get('h2', [])]}\n"
            
            paa_summary = "\n".join([f"- {q}" for q in paa_questions])
            
            system_prompt = """Analyze competitor content and identify content gaps. Return JSON:
            {
                "content_gaps": ["gap1", "gap2", "gap3"],
                "missing_topics": ["topic1", "topic2"],
                "opportunities": ["opportunity1", "opportunity2"]
            }"""
            
            human_prompt = f"""Target: "{title}"
            
            Competitors: {competitor_summary}
            
            PAA Questions: {paa_summary}
            
            Identify content gaps and opportunities to create better content."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            try:
                import json
                return json.loads(response.content)
            except:
                return {"analysis": response.content}
                
        except Exception as e:
            logger.error(f"Error analyzing content gaps: {e}")
            return {"error": str(e)}
    
    def _extract_competitor_keywords(self, scraped_data, title):
        """STEP 4: Extract keywords from competitors using Gemini"""
        if not self.llm:
            return {"error": "Gemini not available"}
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            # Prepare content for analysis
            content_text = ""
            for blog in scraped_data:
                content_text += f"Title: {blog.get('title', '')}\n"
                content_text += f"Meta: {blog.get('metadescription', '')}\n"
                for h in blog.get('h2', []):
                    content_text += f"H2: {h['heading']}\n"
            
            system_prompt = """Extract keywords from competitor content. Return JSON:
            {
                "primary_keywords": ["keyword1", "keyword2"],
                "long_tail_keywords": ["long tail 1", "long tail 2"],
                "secondary_keywords": ["secondary1", "secondary2"]
            }"""
            
            human_prompt = f"""Target: "{title}"
            
            Competitor Content: {content_text[:2000]}
            
            Extract relevant keywords for SEO optimization."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            try:
                import json
                return json.loads(response.content)
            except:
                return {"keywords": response.content}
                
        except Exception as e:
            logger.error(f"Error extracting keywords: {e}")
            return {"error": str(e)}
    
    def _analyze_content_structure(self, scraped_data):
        """STEP 5: Analyze competitor content structure using Gemini"""
        if not self.llm:
            return {"error": "Gemini not available"}
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            # Analyze structure
            structure_info = ""
            for i, blog in enumerate(scraped_data, 1):
                structure_info += f"\nCompetitor {i}:\n"
                structure_info += f"H1 count: {len(blog.get('h1', []))}\n"
                structure_info += f"H2 count: {len(blog.get('h2', []))}\n"
                structure_info += f"H3 count: {len(blog.get('h3', []))}\n"
            
            system_prompt = """Analyze content structure and recommend optimal structure. Return JSON:
            {
                "recommended_structure": {
                    "h1_count": 1,
                    "h2_count": 5,
                    "h3_count": 8,
                    "word_count_range": "2000-3000"
                },
                "content_patterns": ["pattern1", "pattern2"]
            }"""
            
            human_prompt = f"""Competitor Structure Analysis: {structure_info}
            
            Recommend optimal content structure for better SEO."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            try:
                import json
                return json.loads(response.content)
            except:
                return {"structure": response.content}
                
        except Exception as e:
            logger.error(f"Error analyzing structure: {e}")
            return {"error": str(e)}
    
    def _generate_blog_outline(self, title, content_gaps, keywords, structure, target_word_count):
        """STEP 6: Generate detailed blog outline using Gemini"""
        if not self.llm:
            return {"error": "Gemini not available"}
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            # Calculate word distribution based on target
            if target_word_count:
                word_count_text = f"Target: {target_word_count} words"
                # Calculate optimal section distribution
                intro_words = max(150, int(target_word_count * 0.1))  # 10% for intro
                conclusion_words = max(150, int(target_word_count * 0.08))  # 8% for conclusion
                remaining_words = target_word_count - intro_words - conclusion_words
                
                # Estimate 5-8 main sections
                num_sections = min(8, max(5, remaining_words // 300))
                words_per_section = remaining_words // num_sections
                
                section_word_guidance = f"Distribute approximately {words_per_section} words per main section"
            else:
                word_count_text = "Target: 2000-3000 words"
                intro_words = 200
                conclusion_words = 200
                words_per_section = 400
                section_word_guidance = "Distribute approximately 400 words per main section"
            
            system_prompt = f"""Create a detailed, word-count targeted blog outline. Return JSON:
            {{
                "title": "{title}",
                "target_word_count": {target_word_count if target_word_count else 2500},
                "sections": [
                    {{
                        "section_type": "introduction",
                        "heading": "Introduction heading",
                        "target_words": {intro_words},
                        "key_points": ["hook", "problem statement", "preview"],
                        "keywords_to_include": ["primary keyword variations"]
                    }},
                    {{
                        "section_type": "main_content",
                        "heading": "Main section heading",
                        "target_words": {words_per_section},
                        "key_points": ["detailed point 1", "detailed point 2"],
                        "keywords_to_include": ["relevant keywords"]
                    }},
                    {{
                        "section_type": "conclusion",
                        "heading": "Conclusion heading", 
                        "target_words": {conclusion_words},
                        "key_points": ["summary", "call to action"],
                        "keywords_to_include": ["primary keyword"]
                    }}
                ],
                "seo_strategy": {{
                    "primary_keyword_density": "1-2%",
                    "secondary_keywords": ["list of secondary keywords"],
                    "internal_linking_opportunities": ["suggested internal links"]
                }},
                "estimated_total_words": {target_word_count if target_word_count else 2500}
            }}"""
            
            human_prompt = f"""Title: "{title}"
            
            Content Gaps to Address: {content_gaps}
            Keywords to Integrate: {keywords}
            Competitor Structure Analysis: {structure}
            {word_count_text}
            
            Create a comprehensive, word-count optimized blog outline that:
            1. Precisely targets {target_word_count if target_word_count else '2000-3000'} words
            2. {section_word_guidance}
            3. Addresses all identified content gaps
            4. Strategically incorporates extracted keywords
            5. Follows SEO best practices for structure
            6. Ensures each section has clear word count targets that sum to the total target
            7. If Necessary , use tables for comparison and clarity
            8. Include FAQ section if relevant

            Make sure the outline is detailed enough to guide precise content creation."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            try:
                import json
                return json.loads(response.content)
            except:
                return {"outline": response.content}
                
        except Exception as e:
            logger.error(f"Error generating outline: {e}")
            return {"error": str(e)}
    
    def _generate_final_blog(self, title, outline, target_word_count):
        """STEP 7: Generate final SEO optimized blog using Gemini"""
        if not self.llm:
            return f"<h1>{title}</h1>\n<p>Blog generation failed - Gemini not available.</p>"
        
        try:
            from langchain_core.messages import SystemMessage, HumanMessage
            
            # Extract word count target from outline if available
            outline_target = None
            if isinstance(outline, dict):
                outline_target = outline.get('target_word_count') or outline.get('estimated_total_words')
            
            final_target = target_word_count or outline_target or 2500
            
            system_prompt = f"""You are an expert SEO content writer. Write a complete, high-quality blog post in HTML format.

            CRITICAL OUTPUT FORMAT:
            - Output ONLY clean HTML markup that goes inside the <body> tag
            - DO NOT include <body>, </body>, <html>, <head>, or any other wrapper tags
            - Start directly with content (e.g., <h1>, <p>, <ul>, etc.)
            - NO markdown syntax (no **, *, #, etc.) - use proper HTML tags only
            - Use semantic HTML5 tags
            
            CRITICAL CONTENT STRUCTURE:
            
            1. INTRODUCTION (1-2 paragraphs only):
               - Use <h1> for the main title: "{title}"
               - Use <p> tags for paragraphs
               - Brief, engaging hook
               - Quick overview of what readers will learn
            
            2. MAIN CONTENT (Detailed with bullet points):
               - Use <h2> for major sections
               - Use <h3> for subsections
               - Use <ul> and <li> for bullet points
               - Use <ol> and <li> for numbered lists
               - Include structured elements:
                 * Key Features (as <ul> lists)
                 * Pros and Cons (as <ul> lists with clear headings)
                 * Step-by-step instructions (as <ol> lists)
                 * Comparisons (use <table> with inline CSS for mobile responsiveness)
               - Use <strong> ONLY for truly important keywords/phrases (be selective, not everything)
               - Maximum 2-3 sentences per <p> tag
            
            3. TABLES (if needed for comparisons):
               - Include this exact inline CSS in the table tag for mobile responsiveness:
                 <table style="width:100%; border-collapse:collapse; margin:20px 0; display:block; overflow-x:auto; white-space:nowrap;">
               - Use <thead>, <tbody>, <th>, <td> properly
               - Add inline styles: border:1px solid #ddd; padding:12px; text-align:left;
            
            4. HYPERLINKS (for tools/resources):
               - Use <a href="URL" rel="nofollow" target="_blank">Tool Name</a>
               - Always add rel="nofollow" for external tool/resource links
               - Add target="_blank" to open in new tab
            
            5. CONCLUSION (1 paragraph only):
               - Use <h2>Conclusion</h2>
               - Brief summary in <p> tags
               - Call to action if appropriate
            
            6. FAQ SECTION (3-5 questions):
               - Use <h2>Frequently Asked Questions</h2>
               - For each FAQ use: <h3>Question here?</h3> followed by <p>answer</p>
               - Keep answers brief (2-3 sentences max)
            
            TEXT HIGHLIGHTING RULES:
            - Use <strong> SPARINGLY - only for 3-5 most important terms/phrases per section
            - Do NOT highlight every keyword or point
            - Highlight only: key definitions, critical statistics, main takeaways
            - Avoid over-highlighting - it reduces impact
            
            FORMATTING REQUIREMENTS:
            - Target word count: {final_target} words (±50 words)
            - Clean, semantic HTML5
            - No inline CSS except for tables (mobile responsiveness)
            - SEO optimized with strategic keyword placement
            - Professional, scannable format
            - Proper HTML hierarchy
            
            CRITICAL: Return ONLY the HTML content without <body> tags, markdown syntax, or any wrappers.
            The output should be ready to copy-paste directly into a CMS body editor."""
            
            human_prompt = f"""Write a complete blog post with title: "{title}"

            DETAILED OUTLINE TO FOLLOW:
            {outline}
            
            TARGET WORD COUNT: {final_target} words
            
            MANDATORY HTML STRUCTURE:
            
            <h1>{title}</h1>
            
            <p>Introduction paragraph 1...</p>
            <p>Introduction paragraph 2 (if needed)...</p>
            
            <h2>Main Section Title</h2>
            <p>Brief intro to section...</p>
            <ul>
              <li>Point 1 with details</li>
              <li>Point 2 with details</li>
            </ul>
            
            <h3>Subsection if needed</h3>
            <ol>
              <li>Step 1</li>
              <li>Step 2</li>
            </ol>
            
            [If comparison table needed:]
            <table style="width:100%; border-collapse:collapse; margin:20px 0; display:block; overflow-x:auto; white-space:nowrap;">
              <thead>
                <tr>
                  <th style="border:1px solid #ddd; padding:12px; text-align:left;">Header 1</th>
                  <th style="border:1px solid #ddd; padding:12px; text-align:left;">Header 2</th>
                </tr>
              </thead>
              <tbody>
                <tr>
                  <td style="border:1px solid #ddd; padding:12px;">Data 1</td>
                  <td style="border:1px solid #ddd; padding:12px;">Data 2</td>
                </tr>
              </tbody>
            </table>
            
            <h2>Conclusion</h2>
            <p>Summary paragraph...</p>
            
            <h2>Frequently Asked Questions</h2>
            <h3>Question 1?</h3>
            <p>Answer 1...</p>
            <h3>Question 2?</h3>
            <p>Answer 2...</p>
            
            CONTENT GUIDELINES:
            - NO markdown syntax - pure HTML only
            - Integrate all keywords naturally from the outline
            - Add rel="nofollow" to external tool/resource links
            - Use <strong> sparingly for only the most important 3-5 terms per section
            - Make content highly scannable
            - Ensure mobile-responsive tables with inline CSS
            - Professional, authoritative tone
            - Ready to copy-paste into WordPress/CMS
            
            IMPORTANT: Generate the COMPLETE blog post with all sections fully written out. Do not use placeholders, ellipsis (...), or [content here] markers. Write every section in full detail from start to finish.
            
            Begin writing the complete HTML blog post now. Remember: NO <body> tags, NO markdown syntax."""
            
            response = self.llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)])
            
            # Clean up any potential body tags or markdown that might slip through
            content = response.content.strip()
            
            # Remove body tags if present
            import re
            content = re.sub(r'<body[^>]*>|</body>', '', content, flags=re.IGNORECASE)
            content = re.sub(r'<html[^>]*>|</html>', '', content, flags=re.IGNORECASE)
            content = re.sub(r'<head[^>]*>.*?</head>', '', content, flags=re.IGNORECASE | re.DOTALL)
            content = re.sub(r'<!DOCTYPE[^>]*>', '', content, flags=re.IGNORECASE)
            
            # Remove markdown code blocks if present
            content = re.sub(r'```html\n?', '', content)
            content = re.sub(r'```\n?', '', content)
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error generating final blog: {e}")
            return f"<h1>{title}</h1>\n<p>Blog generation failed: {str(e)}</p>"

    def _normalize_model_output_to_markdown(self, raw_text: str) -> str:
        """Normalize model output to clean HTML text"""
        try:
            if not raw_text:
                return ''

            text = raw_text.strip()

            # Remove code fences (html, markdown, etc.)
            if text.startswith('```') and text.endswith('```'):
                parts = text.split('\n')
                if parts and parts[0].startswith('```'):
                    parts = parts[1:]
                if parts and parts[-1].startswith('```'):
                    parts = parts[:-1]
                text = '\n'.join(parts).strip()

            # Remove any remaining code fence markers
            import re
            text = re.sub(r'```html\n?', '', text)
            text = re.sub(r'```\n?', '', text)

            # Try to extract from JSON if present
            if text.startswith('{'):
                try:
                    import json
                    parsed = json.loads(text)
                    if isinstance(parsed, dict):
                        if 'final_article' in parsed:
                            return parsed['final_article'].strip()
                        elif 'blog_post' in parsed and 'content' in parsed['blog_post']:
                            return parsed['blog_post']['content'].strip()
                except:
                    pass

            # Remove body tags if they somehow got included
            text = re.sub(r'<body[^>]*>|</body>', '', text, flags=re.IGNORECASE)
            text = re.sub(r'<html[^>]*>|</html>', '', text, flags=re.IGNORECASE)
            text = re.sub(r'<head[^>]*>.*?</head>', '', text, flags=re.IGNORECASE | re.DOTALL)
            text = re.sub(r'<!DOCTYPE[^>]*>', '', text, flags=re.IGNORECASE)

            return text.strip()
        except:
            return raw_text if isinstance(raw_text, str) else str(raw_text)

    def save_as_document(self, formatted_post, filename):
        """Save the formatted blog post as a .docx file"""
        try:
            from docx import Document
            import re
            
            file_path = os.path.join(self.media_root, 'generated_blogs', filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            doc = Document()
            
            lines = formatted_post.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                else:
                    # Clean markdown formatting
                    clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)  # Remove bold
                    clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)  # Remove italic
                    clean_text = re.sub(r'`([^`]+)`', r'\1', clean_text)  # Remove code
                    doc.add_paragraph(clean_text)
            
            doc.save(file_path)
            logger.info(f"Document saved to: {file_path}")
            return file_path
                
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return None
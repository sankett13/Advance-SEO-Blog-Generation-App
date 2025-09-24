# version 2 for the SEO content automation script
# version 1 : data_scrap_test.py

import requests
from bs4 import BeautifulSoup
import os
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, SystemMessage
from urllib.parse import urlparse
from datetime import datetime
import json
from docx import Document
import re
# from client import RestClient

load_dotenv()

# API Keys and Configuration
SERP_API_KEY = os.getenv("SERP_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_GEMINI_API_KEY")
DATAFORSEO_LOGIN = os.getenv("DATAFORSEO_LOGIN")
DATAFORSEO_PASSWORD = os.getenv("DATAFORSEO_PASSWORD")

# Initialize Gemini 2.5 Pro model
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-pro",
    google_api_key=GOOGLE_API_KEY,
    temperature=0.3,
    max_tokens=8192
)

headings = {
    "h1": [],
    "h2": [],
    "h3": []
}


def get_top_5_links_and_paa(query, api_key, num_results=3):
    if not api_key:
        return [], []

    params = {
        "q": query,
        "api_key": api_key,
        "engine": "google",
        "num": 20,
        "hl": "en",
        "gl": "us",
    }

    # Fetch results (prefer serpapi client if available, fallback to HTTP)
    try:
        try:
            from serpapi import GoogleSearch  # type: ignore
            client = GoogleSearch(params)
            results = client.get_dict()
        except Exception:
            resp = requests.get("https://serpapi.com/search", params=params, timeout=12)
            resp.raise_for_status()
            results = resp.json()
    except Exception:
        return [], []

    organic_results = results.get("organic_results", []) or []
    paa_results = results.get("related_questions", []) or []

    # Exclude specific domains
    exclude_domains = {
        "reddit.com", "www.reddit.com",
        "wikipedia.org", "en.wikipedia.org",
        "quora.com", "www.quora.com",
        "facebook.com", "www.facebook.com",
        "twitter.com", "x.com", "www.twitter.com", "www.x.com",
        "linkedin.com", "www.linkedin.com",
        "pinterest.com", "www.pinterest.com",
        "youtube.com", "www.youtube.com",
        "tiktok.com", "www.tiktok.com",
        "instagram.com", "www.instagram.com",
    }

    selected_links = []
    seen_domains = set()

    for r in organic_results:
        link = r.get("link")
        if not link:
            continue
        domain = urlparse(link).netloc.lower()
        if domain.startswith("www."):
            domain = domain[4:]
        if domain in exclude_domains or domain in seen_domains:
            continue
        seen_domains.add(domain)
        selected_links.append(link)
        if len(selected_links) >= num_results:
            break

    paa_questions = []
    for q in paa_results:
        question = q.get("question")
        if question:
            paa_questions.append(question)

    return selected_links[:num_results], paa_questions


def scrape_blog(url):
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        response = requests.get(url, headers=headers, timeout=10)
    except Exception:
        return None

    soup = BeautifulSoup(response.text, "html.parser")

    # Title
    title = soup.title.string.strip() if soup.title and soup.title.string else "No title"

    # Meta description
    meta_desc_tag = soup.find("meta", attrs={"name": "description"})
    metadescription = meta_desc_tag["content"].strip() if meta_desc_tag and meta_desc_tag.get("content") else ""

    # Tagline
    tagline = ""
    tagline_meta = soup.find("meta", attrs={"name": "tagline"})
    if tagline_meta and tagline_meta.get("content"):
        tagline = tagline_meta["content"].strip()
    else:
        tagline_p = soup.find("p", class_="tagline")
        if tagline_p:
            tagline = tagline_p.get_text(strip=True)

    # Helper to get paragraphs under a heading
    def get_paras_after(tag):
        paras = []
        next_node = tag.find_next_sibling()
        while next_node and next_node.name not in ["h1", "h2", "h3"]:
            if next_node.name == "p":
                paras.append(next_node.get_text(strip=True))
            next_node = next_node.find_next_sibling()
        return paras

    # Headings and their paragraphs
    h1 = []
    for h in soup.find_all("h1"):
        h1.append({
            "heading": h.get_text(strip=True),
            "paras": get_paras_after(h)
        })
    h2 = []
    for h in soup.find_all("h2"):
        h2.append({
            "heading": h.get_text(strip=True),
            "paras": get_paras_after(h)
        })
    h3 = []
    for h in soup.find_all("h3"):
        h3.append({
            "heading": h.get_text(strip=True),
            "paras": get_paras_after(h)
        })

    # Whole body content
    # body_tag = soup.body
    # body_content = body_tag.get_text(separator="\n", strip=True) if body_tag else ""

    return {
        "url": url,
        "title": title,
        "metadescription": metadescription,
        "tagline": tagline,
        "h1": h1,
        "h2": h2,
        "h3": h3,
        # "body_content": body_content
    }


def analyze_content_gaps(scraped_blogs_data, paa_questions, target_keyword):
    """
    Use Gemini to analyze content gaps between competitors and identify opportunities
    
    Args:
        scraped_blogs_data (list): List of scraped blog data
        paa_questions (list): People Also Ask questions
        target_keyword (str): Target keyword for analysis
        
    Returns:
        dict: Content gap analysis results
    """
    print("\n🔍 Analyzing Content Gaps with Gemini...")
    
    # Prepare the analysis prompt
    competitors_summary = ""
    for i, blog in enumerate(scraped_blogs_data, 1):
        if blog:
            competitors_summary += f"\n--- Competitor {i} ({blog['url']}) ---\n"
            competitors_summary += f"Title: {blog['title']}\n"
            competitors_summary += f"Meta Description: {blog['metadescription']}\n"
            competitors_summary += f"H1 Headings: {[h['heading'] for h in blog['h1']]}\n"
            competitors_summary += f"H2 Headings: {[h['heading'] for h in blog['h2']]}\n"
            competitors_summary += f"H3 Headings: {[h['heading'] for h in blog['h3']]}\n"
    
    paa_summary = "\n".join([f"- {q}" for q in paa_questions])
    
    system_prompt = """You are an expert SEO content strategist and competitor analysis specialist. 
    Your task is to analyze competitor blog content and identify content gaps and opportunities.
    
    Provide your analysis in the following JSON format:
    {
        "content_gaps": [
            {"gap": "description of content gap", "opportunity": "how to address this gap", "priority": "high/medium/low"}
        ],
        "missing_topics": ["topic1", "topic2", "topic3"],
        "content_structure_insights": {
            "common_patterns": ["pattern1", "pattern2"],
            "unique_approaches": ["approach1", "approach2"],
            "recommended_structure": ["section1", "section2", "section3"]
        },
        "paa_coverage_analysis": {
            "covered_questions": ["question1", "question2"],
            "missed_questions": ["question1", "question2"],
            "new_question_opportunities": ["question1", "question2"]
        }
    }"""
    
    human_prompt = f"""
    Target Keyword: "{target_keyword}"
    
    Competitor Blog Analysis:
    {competitors_summary}
    
    People Also Ask Questions:
    {paa_summary}
    
    Please analyze:
    1. What content gaps exist between competitors?
    2. Which People Also Ask questions are not being addressed adequately?
    3. What content structure patterns are successful?
    4. What opportunities exist to create better content?
    
    Focus on actionable insights for creating superior content that can outrank competitors.
    """
    
    try:
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        
        response = llm.invoke(messages)
        
        # Try to parse JSON response
        try:
            analysis_result = json.loads(response.content)
        except json.JSONDecodeError:
            # If JSON parsing fails, return the raw content
            analysis_result = {"raw_analysis": response.content}
        
        return analysis_result
        
    except Exception as e:
        print(f"Error in content gap analysis: {e}")
        return {"error": str(e)}


def extract_competitor_keywords(scraped_blogs_data, target_keyword):
    """
    Use Gemini to identify important keywords used by competitors
    
    Args:
        scraped_blogs_data (list): List of scraped blog data
        target_keyword (str): Target keyword for context
        
    Returns:
        dict: Competitor keyword analysis
    """
    print("\n🔍 Extracting Competitor Keywords with Gemini...")
    
    # Prepare competitor content for analysis
    competitor_content = ""
    for i, blog in enumerate(scraped_blogs_data, 1):
        if blog:
            competitor_content += f"\n--- Competitor {i} ---\n"
            competitor_content += f"Title: {blog['title']}\n"
            competitor_content += f"Meta Description: {blog['metadescription']}\n"
            
            # Include all heading content
            for h1 in blog['h1']:
                competitor_content += f"H1: {h1['heading']}\n"
                for para in h1['paras']:
                    competitor_content += f"Para: {para[:200]}...\n"  # Limit paragraph length
            
            for h2 in blog['h2']:
                competitor_content += f"H2: {h2['heading']}\n"
                for para in h2['paras']:
                    competitor_content += f"Para: {para[:200]}...\n"
    
    system_prompt = """You are an expert SEO keyword researcher and content analyst.
    Your task is to identify the most important keywords and phrases used by competitors.
    
    Provide your analysis in the following JSON format:
    {
        "primary_keywords": [
            {"keyword": "keyword phrase", "frequency": "high/medium/low", "context": "where it appears most"}
        ],
        "long_tail_keywords": [
            {"keyword": "long tail phrase", "intent": "informational/commercial/navigational", "ranking_potential": "high/medium/low"}
        ],
        "semantic_keywords": ["related term 1", "related term 2", "related term 3"],
        "competitor_keyword_strategies": {
            "title_optimization": ["common title patterns"],
            "heading_strategies": ["h1 patterns", "h2 patterns"],
            "content_focus": ["main topic clusters"]
        },
        "keyword_opportunities": [
            {"keyword": "missed keyword", "reason": "why it's an opportunity", "difficulty": "estimated difficulty"}
        ]
    }"""
    
    human_prompt = f"""
    Target Keyword Context: "{target_keyword}"
    
    Competitor Content Analysis:
    {competitor_content}
    
    Please analyze and identify:
    1. Primary keywords used across competitor titles, headings, and content
    2. Long-tail keyword phrases that appear frequently
    3. Semantic/related keywords that support the main topic
    4. Keyword strategies used in titles and headings
    5. Keyword opportunities that competitors might be missing
    
    Focus on keywords that would be valuable for SEO ranking and content optimization.
    """
    
    try:
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        
        response = llm.invoke(messages)
        
        try:
            keyword_analysis = json.loads(response.content)
        except json.JSONDecodeError:
            keyword_analysis = {"raw_analysis": response.content}
        
        return keyword_analysis
        
    except Exception as e:
        print(f"Error in keyword extraction: {e}")
        return {"error": str(e)}


def analyze_competitor_structure(scraped_blogs_data, target_keyword):
    """
    Use Gemini to analyze competitor content structure and identify patterns
    
    Args:
        scraped_blogs_data (list): List of scraped blog data
        target_keyword (str): Target keyword for context
        
    Returns:
        dict: Competitor structure analysis
    """
    print("\n🔍 Analyzing Competitor Content Structure with Gemini...")
    
    structure_data = ""
    for i, blog in enumerate(scraped_blogs_data, 1):
        if blog:
            structure_data += f"\n--- Competitor {i} Structure ---\n"
            structure_data += f"URL: {blog['url']}\n"
            structure_data += f"Title: {blog['title']}\n"
            structure_data += f"Meta Description Length: {len(blog['metadescription'])} chars\n"
            structure_data += f"Number of H1s: {len(blog['h1'])}\n"
            structure_data += f"Number of H2s: {len(blog['h2'])}\n"
            structure_data += f"Number of H3s: {len(blog['h3'])}\n"
            
            structure_data += "Content Structure:\n"
            for h1 in blog['h1']:
                structure_data += f"  H1: {h1['heading']} ({len(h1['paras'])} paragraphs)\n"
            for h2 in blog['h2']:
                structure_data += f"  H2: {h2['heading']} ({len(h2['paras'])} paragraphs)\n"
    
    system_prompt = """You are an expert content strategist and SEO specialist.
    Your task is to analyze competitor content structures and identify winning patterns.
    
    Provide your analysis in the following JSON format:
    {
        "structure_patterns": {
            "title_length_range": "X-Y characters",
            "meta_description_patterns": ["pattern1", "pattern2"],
            "heading_hierarchy": ["typical h1/h2/h3 structure"],
            "content_length_insights": "insights about content depth"
        },
        "successful_elements": [
            {"element": "structure element", "benefit": "why it works", "frequency": "how often used"}
        ],
        "content_organization": {
            "intro_patterns": ["how competitors start content"],
            "main_section_patterns": ["how they organize main content"],
            "conclusion_patterns": ["how they conclude"]
        },
        "optimization_opportunities": [
            {"opportunity": "what could be improved", "implementation": "how to implement", "impact": "expected impact"}
        ],
        "recommended_structure": {
            "title_formula": "recommended title approach",
            "meta_description": "recommended meta description approach",
            "heading_structure": ["recommended H1", "recommended H2s", "recommended H3s"],
            "content_sections": ["section 1", "section 2", "section 3"]
        }
    }"""
    
    human_prompt = f"""
    Target Keyword: "{target_keyword}"
    
    Competitor Structure Analysis:
    {structure_data}
    
    Please analyze:
    1. Common structural patterns across top-ranking competitors
    2. Title and meta description optimization strategies
    3. Heading hierarchy and organization patterns
    4. Content depth and organization approaches
    5. Opportunities to create a superior content structure
    
    Focus on actionable insights for creating content that can outrank competitors through better structure and organization.
    """
    
    try:
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        
        response = llm.invoke(messages)
        
        try:
            structure_analysis = json.loads(response.content)
        except json.JSONDecodeError:
            structure_analysis = {"raw_analysis": response.content}
        
        return structure_analysis
        
    except Exception as e:
        print(f"Error in structure analysis: {e}")
        return {"error": str(e)}


def comprehensive_competitor_analysis(target_keyword, num_results=3):
    """
    Run a comprehensive competitor analysis including scraping, content gaps, keywords, and structure
    
    Args:
        target_keyword (str): Target keyword to analyze
        num_results (int): Number of competitors to analyze
        
    Returns:
        dict: Complete analysis results
    """
    print(f"\n🚀 Starting Comprehensive Competitor Analysis for: '{target_keyword}'")
    print("=" * 80)
    
    # Step 1: Get top competitors and PAA questions
    print("\n1. 🔍 Getting top competitors and People Also Ask questions...")
    links, paa_questions = get_top_5_links_and_paa(target_keyword, SERP_API_KEY, num_results)
    
    if not links:
        print("❌ No competitor links found. Please check your SERP_API_KEY.")
        return {"error": "No competitor data available"}
    
    print(f"✅ Found {len(links)} competitors and {len(paa_questions)} PAA questions")
    
    # Step 2: Scrape competitor blogs
    print("\n2. 📊 Scraping competitor blog content...")
    scraped_data = []
    for i, link in enumerate(links, 1):
        print(f"   Scraping competitor {i}: {link}")
        blog_data = scrape_blog(link)
        if blog_data:
            scraped_data.append(blog_data)
            print(f"   ✅ Successfully scraped competitor {i}")
        else:
            print(f"   ❌ Failed to scrape competitor {i}")
    
    if not scraped_data:
        print("❌ No blog data could be scraped.")
        return {"error": "No blog data available"}
    
    # Step 3: AI-powered analysis
    print(f"\n3. 🤖 Running AI analysis with Gemini 2.5 Flash...")
    
    analysis_results = {
        "target_keyword": target_keyword,
        "competitors_analyzed": len(scraped_data),
        "paa_questions": paa_questions,
        "scraped_data": scraped_data,
        "content_gaps": None,
        "competitor_keywords": None,
        "content_structure": None
    }
    
    # Content Gap Analysis
    print("\n   📋 Analyzing content gaps...")
    content_gaps = analyze_content_gaps(scraped_data, paa_questions, target_keyword)
    analysis_results["content_gaps"] = content_gaps
    
    # Keyword Extraction
    print("\n   🔑 Extracting competitor keywords...")
    competitor_keywords = extract_competitor_keywords(scraped_data, target_keyword)
    analysis_results["competitor_keywords"] = competitor_keywords
    
    # Structure Analysis
    print("\n   🏗️  Analyzing content structure...")
    structure_analysis = analyze_competitor_structure(scraped_data, target_keyword)
    analysis_results["content_structure"] = structure_analysis
    
    return analysis_results


def display_analysis_results(analysis_results):
    """
    Display the comprehensive analysis results in a formatted way
    
    Args:
        analysis_results (dict): Results from comprehensive_competitor_analysis
    """
    if "error" in analysis_results:
        print(f"❌ Error: {analysis_results['error']}")
        return
    
    target_keyword = analysis_results["target_keyword"]
    print(f"\n📊 COMPREHENSIVE COMPETITOR ANALYSIS RESULTS")
    print(f"🎯 Target Keyword: '{target_keyword}'")
    print(f"👥 Competitors Analyzed: {analysis_results['competitors_analyzed']}")
    print("=" * 80)
    
    # Content Gaps Summary
    if analysis_results.get("content_gaps") and "content_gaps" in analysis_results["content_gaps"]:
        print(f"\n🔍 CONTENT GAPS IDENTIFIED:")
        gaps = analysis_results["content_gaps"]["content_gaps"]
        for i, gap in enumerate(gaps[:5], 1):  # Show top 5 gaps
            print(f"   {i}. {gap.get('gap', 'N/A')} (Priority: {gap.get('priority', 'N/A')})")
    
    # Top Keywords Summary
    if analysis_results.get("competitor_keywords") and "primary_keywords" in analysis_results["competitor_keywords"]:
        print(f"\n🔑 TOP COMPETITOR KEYWORDS:")
        keywords = analysis_results["competitor_keywords"]["primary_keywords"]
        for i, kw in enumerate(keywords[:10], 1):  # Show top 10 keywords
            print(f"   {i}. '{kw.get('keyword', 'N/A')}' (Frequency: {kw.get('frequency', 'N/A')})")
    
    # Structure Insights Summary
    if analysis_results.get("content_structure") and "recommended_structure" in analysis_results["content_structure"]:
        print(f"\n🏗️  RECOMMENDED CONTENT STRUCTURE:")
        rec_structure = analysis_results["content_structure"]["recommended_structure"]
        print(f"   Title: {rec_structure.get('title_formula', 'N/A')}")
        print(f"   Meta: {rec_structure.get('meta_description', 'N/A')}")
        if rec_structure.get('content_sections'):
            print(f"   Sections:")
            for section in rec_structure['content_sections']:
                print(f"     - {section}")
    
    # PAA Questions
    print(f"\n❓ PEOPLE ALSO ASK QUESTIONS ({len(analysis_results['paa_questions'])}):")
    for i, question in enumerate(analysis_results['paa_questions'][:5], 1):
        print(f"   {i}. {question}")
    
    print(f"\n✅ Analysis Complete! Full results saved to JSON file.")


def generate_blog_outline(analysis_results):
    """
    Generate comprehensive blog outline and structure based on competitor analysis
    
    Args:
        analysis_results (dict): Results from comprehensive_competitor_analysis
        
    Returns:
        dict: Complete blog outline with title, headings, meta description, etc.
    """
    print("\n🎯 Generating Optimized Blog Outline with Gemini...")
    
    target_keyword = analysis_results.get("target_keyword", "")
    paa_questions = analysis_results.get("paa_questions", [])
    
    # Extract content gaps information
    content_gaps_text = ""
    if analysis_results.get("content_gaps"):
        content_gaps_text = json.dumps(analysis_results["content_gaps"], indent=2)
    
    # Extract competitor keywords information  
    keywords_text = ""
    if analysis_results.get("competitor_keywords"):
        keywords_text = json.dumps(analysis_results["competitor_keywords"], indent=2)
    
    # Extract content structure information
    structure_text = ""
    if analysis_results.get("content_structure"):
        structure_text = json.dumps(analysis_results["content_structure"], indent=2)
    
    # Analyze competitor titles and meta descriptions for patterns
    competitor_titles = []
    competitor_metas = []
    competitor_structures = []
    
    for competitor in analysis_results.get("scraped_data", []):
        if competitor:
            competitor_titles.append(competitor.get("title", ""))
            competitor_metas.append(competitor.get("metadescription", ""))
            
            # Structure analysis
            h1_count = len(competitor.get("h1", []))
            h2_count = len(competitor.get("h2", []))
            h3_count = len(competitor.get("h3", []))
            competitor_structures.append({
                "url": competitor.get("url", ""),
                "h1_count": h1_count,
                "h2_count": h2_count,
                "h3_count": h3_count,
                "h2_headings": [h["heading"] for h in competitor.get("h2", [])]
            })
    
    system_prompt = """You are an expert SEO content strategist and blog outline creator. Your task is to create a comprehensive, SEO-optimized blog outline that will outrank competitors while fulfilling user intent and following E-E-A-T (Experience, Expertise, Authoritativeness, Trustworthiness) guidelines.

Your response MUST be in valid JSON format with this exact structure:
{
    "blog_outline": {
        "title": "SEO-optimized title (50-60 characters)",
        "meta_description": "Compelling meta description (150-160 characters)",
        "tagline": "Engaging tagline/subtitle",
        "h1": "Main heading (can be same as title or slightly different)",
        "content_structure": {
            "estimated_word_count": "2500-3500 words",
            "reading_time": "8-12 minutes",
            "target_audience": "description of target audience"
        },
        "heading_hierarchy": [
            {
                "type": "h2",
                "heading": "Introduction heading",
                "purpose": "Hook readers and establish context",
                "word_count": "200-300",
                "key_points": ["point 1", "point 2", "point 3"],
                "keywords_to_include": ["keyword 1", "keyword 2"]
            },
            {
                "type": "h2", 
                "heading": "Main section heading",
                "purpose": "Core content purpose",
                "word_count": "400-600",
                "subsections": [
                    {
                        "type": "h3",
                        "heading": "Subsection heading",
                        "purpose": "Specific purpose",
                        "content_elements": ["comparison table", "pros/cons", "screenshots"]
                    }
                ],
                "keywords_to_include": ["keyword 1", "keyword 2"]
            }
        ],
        "eeat_framework": {
            "experience": ["How to demonstrate experience"],
            "expertise": ["How to show expertise"], 
            "authoritativeness": ["How to establish authority"],
            "trustworthiness": ["How to build trust"]
        },
        "content_gaps_addressed": ["gap 1", "gap 2", "gap 3"],
        "paa_questions_coverage": [
            {"question": "PAA question", "section": "which section addresses it"}
        ],
        "keyword_strategy": {
            "primary_keyword": "main target keyword",
            "secondary_keywords": ["keyword 1", "keyword 2", "keyword 3"],
            "long_tail_keywords": ["long tail 1", "long tail 2"],
            "keyword_density_target": "1-2% for primary, 0.5-1% for secondary"
        },
        "cta_strategy": {
            "primary_cta": "main call to action",
            "secondary_ctas": ["cta 1", "cta 2"],
            "placement": "where to place CTAs"
        }
    }
}

Focus on creating content that:
1. Directly answers user intent better than competitors
2. Naturally incorporates keywords without keyword stuffing
3. Follows E-E-A-T guidelines
4. Addresses content gaps identified in the analysis
5. Has a logical, user-friendly structure
6. Includes actionable insights and value"""

    human_prompt = f"""
Target Keyword: "{target_keyword}"

People Also Ask Questions:
{chr(10).join([f"- {q}" for q in paa_questions])}

Content Gaps Analysis:
{content_gaps_text}

Competitor Keywords Analysis: 
{keywords_text}

Content Structure Analysis:
{structure_text}

Competitor Titles for Reference:
{chr(10).join([f"- {title}" for title in competitor_titles if title])}

Competitor Meta Descriptions for Reference:
{chr(10).join([f"- {meta}" for meta in competitor_metas if meta])}

Competitor Structure Patterns:
{json.dumps(competitor_structures, indent=2)}

Based on this comprehensive analysis, create a blog outline that:

1. **Title & Meta**: Create a compelling title and meta description that outperforms competitors
2. **Structure**: Design a logical heading hierarchy that covers all important aspects
3. **Content Gaps**: Address the identified content gaps to provide superior value
4. **Keywords**: Naturally integrate target keywords and related terms
5. **E-E-A-T**: Include elements that demonstrate experience, expertise, authority, and trust
6. **User Intent**: Ensure the outline fulfills the user's search intent better than competitors
7. **PAA Coverage**: Address People Also Ask questions within the content structure

The goal is to create content that ranks #1 by providing the most comprehensive, valuable, and user-focused resource on this topic.
"""

    try:
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        
        response = llm.invoke(messages)
        
        # Try to parse JSON response
        try:
            # Clean up the response to extract JSON
            content = response.content.strip()
            if content.startswith("```json"):
                content = content[7:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()
            
            blog_outline = json.loads(content)
            return blog_outline
            
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON response: {e}")
            return {"error": "JSON parsing failed", "raw_response": response.content}
        
    except Exception as e:
        print(f"Error generating blog outline: {e}")
        return {"error": str(e)}


def display_blog_outline(outline_data):
    """
    Display the generated blog outline in a formatted, readable way
    
    Args:
        outline_data (dict): Blog outline data from generate_blog_outline
    """
    if "error" in outline_data:
        print(f"❌ Error generating outline: {outline_data['error']}")
        if "raw_response" in outline_data:
            print(f"Raw response: {outline_data['raw_response'][:500]}...")
        return
    
    blog_outline = outline_data.get("blog_outline", {})
    
    print("\n" + "="*80)
    print("🎯 OPTIMIZED BLOG OUTLINE & STRUCTURE")
    print("="*80)
    
    # Basic Information
    print(f"\n📰 TITLE: {blog_outline.get('title', 'N/A')}")
    print(f"📝 META DESCRIPTION: {blog_outline.get('meta_description', 'N/A')}")
    print(f"🏷️  TAGLINE: {blog_outline.get('tagline', 'N/A')}")
    print(f"🎯 H1: {blog_outline.get('h1', 'N/A')}")
    
    # Content Structure Info
    content_structure = blog_outline.get('content_structure', {})
    print(f"\n📊 CONTENT SPECS:")
    print(f"   📏 Word Count: {content_structure.get('estimated_word_count', 'N/A')}")
    print(f"   ⏱️  Reading Time: {content_structure.get('reading_time', 'N/A')}")
    print(f"   👥 Target Audience: {content_structure.get('target_audience', 'N/A')}")
    
    # Heading Hierarchy
    print(f"\n🏗️  CONTENT STRUCTURE:")
    heading_hierarchy = blog_outline.get('heading_hierarchy', [])
    for i, section in enumerate(heading_hierarchy, 1):
        heading_type = section.get('type', 'h2').upper()
        heading_text = section.get('heading', 'N/A')
        purpose = section.get('purpose', 'N/A')
        word_count = section.get('word_count', 'N/A')
        
        print(f"\n   {i}. {heading_type}: {heading_text}")
        print(f"      Purpose: {purpose}")
        print(f"      Word Count: {word_count}")
        
        # Key points
        key_points = section.get('key_points', [])
        if key_points:
            print(f"      Key Points:")
            for point in key_points:
                print(f"        • {point}")
        
        # Keywords to include
        keywords = section.get('keywords_to_include', [])
        if keywords:
            print(f"      Keywords: {', '.join(keywords)}")
        
        # Subsections
        subsections = section.get('subsections', [])
        if subsections:
            print(f"      Subsections:")
            for sub in subsections:
                sub_type = sub.get('type', 'h3').upper()
                sub_heading = sub.get('heading', 'N/A')
                sub_purpose = sub.get('purpose', 'N/A')
                print(f"        {sub_type}: {sub_heading}")
                print(f"          Purpose: {sub_purpose}")
                
                content_elements = sub.get('content_elements', [])
                if content_elements:
                    print(f"          Elements: {', '.join(content_elements)}")
    
    # E-E-A-T Framework
    eeat = blog_outline.get('eeat_framework', {})
    if eeat:
        print(f"\n🏆 E-E-A-T FRAMEWORK:")
        for key, values in eeat.items():
            print(f"   {key.upper()}:")
            for value in values:
                print(f"     • {value}")
    
    # Content Gaps Addressed
    gaps_addressed = blog_outline.get('content_gaps_addressed', [])
    if gaps_addressed:
        print(f"\n🎯 CONTENT GAPS ADDRESSED:")
        for i, gap in enumerate(gaps_addressed, 1):
            print(f"   {i}. {gap}")
    
    # PAA Questions Coverage
    paa_coverage = blog_outline.get('paa_questions_coverage', [])
    if paa_coverage:
        print(f"\n❓ PAA QUESTIONS COVERAGE:")
        for item in paa_coverage:
            question = item.get('question', 'N/A')
            section = item.get('section', 'N/A')
            print(f"   Q: {question}")
            print(f"   A: Covered in {section}")
    
    # Keyword Strategy
    keyword_strategy = blog_outline.get('keyword_strategy', {})
    if keyword_strategy:
        print(f"\n🔑 KEYWORD STRATEGY:")
        print(f"   Primary: {keyword_strategy.get('primary_keyword', 'N/A')}")
        secondary = keyword_strategy.get('secondary_keywords', [])
        if secondary:
            print(f"   Secondary: {', '.join(secondary)}")
        long_tail = keyword_strategy.get('long_tail_keywords', [])
        if long_tail:
            print(f"   Long-tail: {', '.join(long_tail)}")
        print(f"   Density Target: {keyword_strategy.get('keyword_density_target', 'N/A')}")
    
    # CTA Strategy
    cta_strategy = blog_outline.get('cta_strategy', {})
    if cta_strategy:
        print(f"\n📢 CTA STRATEGY:")
        print(f"   Primary CTA: {cta_strategy.get('primary_cta', 'N/A')}")
        secondary_ctas = cta_strategy.get('secondary_ctas', [])
        if secondary_ctas:
            print(f"   Secondary CTAs: {', '.join(secondary_ctas)}")
        print(f"   Placement: {cta_strategy.get('placement', 'N/A')}")
    
    print(f"\n✅ Blog outline generation complete!")


def create_complete_content_strategy(target_keyword, num_competitors=3):
    """
    Complete workflow: Competitor Analysis → Blog Outline Generation
    
    Args:
        target_keyword (str): Target keyword to analyze
        num_competitors (int): Number of competitors to analyze
        
    Returns:
        dict: Complete content strategy including analysis and outline
    """
    print(f"\n🚀 CREATING COMPLETE CONTENT STRATEGY")
    print(f"🎯 Target Keyword: '{target_keyword}'")
    print("="*80)
    
    # Step 1: Competitor Analysis
    print(f"\n📊 STEP 1: COMPETITOR ANALYSIS")
    analysis_results = comprehensive_competitor_analysis(target_keyword, num_competitors)
    
    if "error" in analysis_results:
        return analysis_results
    
    # Step 2: Blog Outline Generation
    print(f"\n📝 STEP 2: BLOG OUTLINE GENERATION")
    blog_outline = generate_blog_outline(analysis_results)
    
    # Combine results
    complete_strategy = {
        "target_keyword": target_keyword,
        "analysis_date": datetime.now().isoformat(),
        "competitor_analysis": analysis_results,
        "blog_outline": blog_outline,
        "strategy_summary": {
            "competitors_analyzed": analysis_results.get("competitors_analyzed", 0),
            "paa_questions_count": len(analysis_results.get("paa_questions", [])),
            "content_gaps_identified": len(analysis_results.get("content_gaps", {}).get("content_gaps", [])) if isinstance(analysis_results.get("content_gaps"), dict) else 0,
            "outline_generated": "error" not in blog_outline
        }
    }
    
    return complete_strategy


def generate_blog_post(blog_outline, analysis_results=None):
    """
    Generate a complete SEO-optimized blog post from the blog outline
    
    Args:
        blog_outline (dict): Blog outline from generate_blog_outline()
        analysis_results (dict): Optional competitor analysis for additional context
        
    Returns:
        dict: Complete blog post with all sections
    """
    print("\n📝 Generating Complete Blog Post with Gemini...")
    
    if "error" in blog_outline:
        return {"error": "Cannot generate blog post due to outline errors"}
    
    # Handle both old and new data structures
    outline_data = blog_outline.get("blog_outline", blog_outline)
    if not outline_data or not isinstance(outline_data, dict):
        return {"error": "Invalid blog outline structure"}
    
    # Check if we have required outline elements
    if not outline_data.get("title") and not outline_data.get("heading_hierarchy"):
        return {"error": "Blog outline missing required elements (title or heading_hierarchy)"}
    
    # Extract key information from outline
    title = outline_data.get("title", "")
    target_keyword = outline_data.get("keyword_strategy", {}).get("primary_keyword", "")
    heading_hierarchy = outline_data.get("heading_hierarchy", [])
    eeat_framework = outline_data.get("eeat_framework", {})
    content_gaps = outline_data.get("content_gaps_addressed", [])
    paa_coverage = outline_data.get("paa_questions_coverage", [])
    
    # Get competitor data for context (if available)
    competitor_context = ""
    if analysis_results and analysis_results.get("scraped_data"):
        competitor_context = "\n\nCompetitor Analysis Context:\n"
        for i, competitor in enumerate(analysis_results["scraped_data"][:3], 1):
            if competitor:
                competitor_context += f"Competitor {i}: {competitor.get('title', 'N/A')}\n"
                competitor_context += f"Key H2s: {[h['heading'] for h in competitor.get('h2', [])]}\n"
    
    system_prompt = """You are an expert SEO content writer and digital marketing specialist. Your task is to write a complete, high-quality blog post that will rank #1 on Google while providing exceptional value to readers.

CRITICAL REQUIREMENTS:
1. **User Intent Focus**: Write ONLY relevant content that directly answers user intent
2. **No Fluff**: Eliminate irrelevant information, generic statements, or filler content
3. **SEO Optimization**: Naturally integrate keywords without keyword stuffing
4. **E-E-A-T Compliance**: Demonstrate experience, expertise, authority, and trustworthiness
5. **Actionable Value**: Provide specific, actionable insights readers can implement
6. **Conversational Tone**: Write in a helpful, authoritative but approachable tone

CONTENT STRUCTURE:
- Write engaging, informative content for each section
- Use specific examples, features, and benefits
- Include practical advice and implementation tips
- Address pain points and provide solutions
- Make comparisons clear and helpful for decision-making

STYLE GUIDELINES:
- Use active voice and clear, concise sentences
- Include specific details (features, pricing, benefits)
- Add transition sentences between sections
- Use bullet points and lists for readability
- Include relevant statistics or data points when appropriate
- Avoid excessive adjectives and marketing fluff

IMPORTANT JSON FORMATTING RULES:
- Escape all quotes inside strings using \"
- Do NOT use unescaped quotes or special characters
- Keep JSON structure simple and clean
- Each string value should be properly escaped
- Use \\n for line breaks within strings if needed

Your response MUST be valid JSON with this exact structure:
{
    "blog_post": {
        "title": "final optimized title",
        "meta_description": "final meta description",
        "content": {
            "introduction": "engaging introduction paragraph(s) - properly escaped",
            "main_sections": [
                {
                    "heading": "section heading",
                    "content": "complete section content with proper formatting - properly escaped",
                    "subsections": [
                        {
                            "subheading": "subsection heading",
                            "content": "subsection content - properly escaped"
                        }
                    ]
                }
            ],
            "conclusion": "compelling conclusion with clear next steps - properly escaped"
        },
        "word_count": "estimated word count",
        "readability_score": "estimated reading level",
        "seo_elements": {
            "primary_keywords_used": ["list of primary keywords naturally integrated"],
            "secondary_keywords_used": ["list of secondary keywords used"],
            "internal_link_opportunities": ["suggested internal link anchor texts"],
            "external_link_opportunities": ["suggested authoritative sources to link to"]
        }
    }
}"""

    human_prompt = f"""
Blog Outline to Convert to Full Post:
{json.dumps(outline_data, indent=2)}

Target Keyword Focus: "{target_keyword}"

Content Gaps to Address:
{json.dumps(content_gaps, indent=2)}

PAA Questions to Cover:
{json.dumps(paa_coverage, indent=2)}

E-E-A-T Framework Guidelines:
{json.dumps(eeat_framework, indent=2)}

{competitor_context}

WRITING INSTRUCTIONS:
1. **Focus on User Intent**: The user wants to find the "top 5 AI tools for email marketing" - provide exactly that with clear comparisons and recommendations
2. **Be Specific**: Include actual features, pricing, use cases, and benefits for each tool
3. **Eliminate Fluff**: No generic statements like "AI is changing the world" - get straight to valuable information
4. **Actionable Content**: Help readers make informed decisions about which tool to choose
5. **Natural Keyword Usage**: Integrate keywords naturally within valuable content
6. **Expert Voice**: Write with authority based on the E-E-A-T guidelines provided

Write a complete blog post that follows the outline structure but focuses entirely on delivering value to someone specifically looking for the top 5 AI email marketing tools. Make every sentence count toward helping them make the best decision for their needs.Stick to the main content and shorten the lenght of the less important parts like introduction and focus and provide more information about the main sections.Also make sure to include the external links where necessary for example to the tools websites or to relevant articles.
"""

    try:
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        
        response = llm.invoke(messages)
        
        # Parse JSON response with improved error handling
        try:
            content = response.content.strip()
            
            # Remove markdown code blocks if present
            if content.startswith("```json"):
                content = content[7:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()
            
            # Try to fix common JSON issues
            # Replace curly quotes with straight quotes
            content = content.replace('"', '"').replace('"', '"')
            content = content.replace(''', "'").replace(''', "'")
            
            # Try parsing the JSON
            blog_post = json.loads(content)
            return blog_post
            
        except json.JSONDecodeError as e:
            print(f"JSON parsing failed: {e}")
            print("Attempting to extract content manually...")
            
            # Fallback: Create a structured response from the raw content
            raw_content = response.content
            
            # Try to extract meaningful sections even if JSON parsing fails
            fallback_post = {
                "blog_post": {
                    "title": "Top 5 AI Email Marketing Tools That Actually Boost Conversions (2025)",
                    "meta_description": "Discover the 5 best AI email marketing tools that increase engagement. Compare features, pricing, and real results from leading platforms.",
                    "content": {
                        "introduction": "",
                        "main_sections": [],
                        "conclusion": ""
                    },
                    "word_count": "2500-3000 words",
                    "seo_elements": {
                        "primary_keywords_used": ["ai email marketing tools", "email marketing", "ai tools"],
                        "secondary_keywords_used": ["email automation", "marketing automation", "ai email"],
                        "internal_link_opportunities": [],
                        "external_link_opportunities": []
                    }
                },
                "raw_response": raw_content,
                "parsing_method": "fallback_structure"
            }
            
            # Try to extract content sections from raw response
            try:
                # Split content into sections and clean it up
                content_sections = raw_content.split('\n\n')
                
                # Simple content extraction
                for i, section in enumerate(content_sections):
                    if section.strip():
                        if i == 0:
                            fallback_post["blog_post"]["content"]["introduction"] = section.strip()
                        elif i == len(content_sections) - 1:
                            fallback_post["blog_post"]["content"]["conclusion"] = section.strip()
                        else:
                            # Create a main section
                            section_data = {
                                "heading": f"Section {i}",
                                "content": section.strip(),
                                "subsections": []
                            }
                            fallback_post["blog_post"]["content"]["main_sections"].append(section_data)
                
            except Exception as extraction_error:
                print(f"Content extraction also failed: {extraction_error}")
                fallback_post["blog_post"]["content"]["introduction"] = raw_content[:1000] + "..."
            
            return fallback_post
        
    except Exception as e:
        print(f"Error generating blog post: {e}")
        print("Attempting alternative blog generation approach...")
        
        # Alternative approach: Generate blog post without strict JSON formatting
        return generate_blog_post_alternative(blog_outline, analysis_results)


def generate_blog_post_alternative(blog_outline, analysis_results=None):
    """
    Alternative blog post generation with simpler prompting (fallback method)
    
    Args:
        blog_outline (dict): Blog outline from generate_blog_outline()
        analysis_results (dict): Optional competitor analysis for additional context
        
    Returns:
        dict: Blog post data in simplified format
    """
    print("\n📝 Using alternative blog generation approach...")
    
    if "error" in blog_outline:
        return {"error": "Cannot generate blog post due to outline errors"}
    
    # Handle both old and new data structures
    outline_data = blog_outline.get("blog_outline", blog_outline)
    if not outline_data or not isinstance(outline_data, dict):
        return {"error": "Invalid blog outline structure"}
    
    # Check if we have required outline elements  
    if not outline_data.get("title") and not outline_data.get("heading_hierarchy"):
        return {"error": "Blog outline missing required elements (title or heading_hierarchy)"}
    
    # Extract key information
    title = outline_data.get("title", "Top 5 AI Email Marketing Tools That Actually Boost Conversions (2025)")
    target_keyword = outline_data.get("keyword_strategy", {}).get("primary_keyword", "ai email marketing tools")
    heading_hierarchy = outline_data.get("heading_hierarchy", [])
    
    # Simplified system prompt for direct content generation
    system_prompt = """You are an expert content writer. Write a comprehensive blog post about the top 5 AI email marketing tools. 

Structure your response as a complete blog post with:
1. An engaging introduction
2. Detailed sections for each of the 5 tools
3. A helpful conclusion

Focus on providing specific, actionable information that helps readers choose the right tool for their needs. Include features, pricing, and use cases for each tool.

Write in a professional but approachable tone. Make sure the content is SEO-friendly and valuable to readers."""
    
    human_prompt = f"""
Write a complete blog post with the title: "{title}"

Target keyword: "{target_keyword}"

The blog post should cover the top 5 AI email marketing tools with detailed analysis of each tool including:
- Key features and capabilities
- Pricing information
- Best use cases
- Pros and cons
- Who should use each tool

Make the content informative, well-structured, and helpful for someone trying to choose the best AI email marketing tool for their business.

Write the complete blog post now:
"""
    
    try:
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        
        response = llm.invoke(messages)
        
        # Create a structured response from the raw content
        blog_post_data = {
            "blog_post": {
                "title": title,
                "meta_description": f"Discover the 5 best {target_keyword} that increase engagement. Compare features, pricing, and real results from leading platforms.",
                "content": {
                    "introduction": "",
                    "main_sections": [
                        {
                            "heading": "Complete AI Email Marketing Tools Analysis",
                            "content": response.content,
                            "subsections": []
                        }
                    ],
                    "conclusion": ""
                },
                "word_count": f"{len(response.content.split())} words",
                "seo_elements": {
                    "primary_keywords_used": [target_keyword, "email marketing", "ai tools"],
                    "secondary_keywords_used": ["email automation", "marketing automation"],
                    "internal_link_opportunities": ["email marketing guide", "ai tools comparison"],
                    "external_link_opportunities": ["tool websites", "industry reports"]
                }
            },
            "generation_method": "alternative_approach"
        }
        
        return blog_post_data
        
    except Exception as e:
        print(f"Alternative blog generation also failed: {e}")
        return {"error": f"Both blog generation methods failed: {str(e)}"}


def format_blog_post_for_publication(blog_post_data):
    """
    Format the generated blog post into a publication-ready format
    
    Args:
        blog_post_data (dict): Blog post data from generate_blog_post()
        
    Returns:
        str: Formatted blog post ready for publication
    """
    if "error" in blog_post_data:
        return f"Error: {blog_post_data['error']}"
    
    blog_post = blog_post_data.get("blog_post", {})
    content = blog_post.get("content", {})
    
    # Start building the formatted post
    formatted_post = ""
    
    # Title
    title = blog_post.get("title", "")
    formatted_post += f"# {title}\n\n"
    
    # Meta description (as comment for reference)
    meta_desc = blog_post.get("meta_description", "")
    formatted_post += f"<!-- Meta Description: {meta_desc} -->\n\n"
    
    # Check if this is from alternative generation method
    if blog_post_data.get("generation_method") == "alternative_approach":
        # For alternative approach, the content is already formatted
        main_sections = content.get("main_sections", [])
        if main_sections and main_sections[0].get("content"):
            formatted_post += main_sections[0]["content"]
        return formatted_post
    
    # Handle fallback structure with raw response
    if blog_post_data.get("parsing_method") == "fallback_structure":
        formatted_post += "## Introduction\n\n"
        intro = content.get("introduction", "")
        if intro:
            formatted_post += f"{intro}\n\n"
        
        # Add main sections
        main_sections = content.get("main_sections", [])
        for section in main_sections:
            heading = section.get("heading", "")
            section_content = section.get("content", "")
            if heading and section_content:
                formatted_post += f"## {heading}\n\n{section_content}\n\n"
        
        conclusion = content.get("conclusion", "")
        if conclusion:
            formatted_post += f"## Conclusion\n\n{conclusion}\n\n"
        
        return formatted_post
    
    # Standard format handling
    # Introduction
    intro = content.get("introduction", "")
    if intro:
        formatted_post += f"{intro}\n\n"
    
    # Main sections
    main_sections = content.get("main_sections", [])
    for section in main_sections:
        # Section heading
        heading = section.get("heading", "")
        if heading:
            formatted_post += f"## {heading}\n\n"
        
        # Section content
        section_content = section.get("content", "")
        if section_content:
            formatted_post += f"{section_content}\n\n"
        
        # Subsections
        subsections = section.get("subsections", [])
        for subsection in subsections:
            subheading = subsection.get("subheading", "")
            if subheading:
                formatted_post += f"### {subheading}\n\n"
            
            subcontent = subsection.get("content", "")
            if subcontent:
                formatted_post += f"{subcontent}\n\n"
    
    # Conclusion
    conclusion = content.get("conclusion", "")
    if conclusion:
        formatted_post += f"## Conclusion\n\n{conclusion}\n\n"
    
    # SEO metadata (as comments)
    seo_elements = blog_post.get("seo_elements", {})
    if seo_elements:
        formatted_post += f"<!-- SEO Information:\n"
        formatted_post += f"Word Count: {blog_post.get('word_count', 'N/A')}\n"
        formatted_post += f"Primary Keywords: {', '.join(seo_elements.get('primary_keywords_used', []))}\n"
        formatted_post += f"Secondary Keywords: {', '.join(seo_elements.get('secondary_keywords_used', []))}\n"
        formatted_post += f"-->\n"
    
    return formatted_post


def display_blog_post_summary(blog_post_data):
    """
    Display a summary of the generated blog post
    
    Args:
        blog_post_data (dict): Blog post data from generate_blog_post()
    """
    if "error" in blog_post_data:
        print(f"❌ Error generating blog post: {blog_post_data['error']}")
        return
    
    blog_post = blog_post_data.get("blog_post", {})
    
    print("\n" + "="*80)
    print("📝 GENERATED BLOG POST SUMMARY")
    print("="*80)
    
    # Basic info
    print(f"\n📰 TITLE: {blog_post.get('title', 'N/A')}")
    print(f"📝 META DESCRIPTION: {blog_post.get('meta_description', 'N/A')}")
    print(f"📊 ESTIMATED WORD COUNT: {blog_post.get('word_count', 'N/A')}")
    print(f"📖 READABILITY: {blog_post.get('readability_score', 'N/A')}")
    
    # Content structure
    content = blog_post.get("content", {})
    main_sections = content.get("main_sections", [])
    
    print(f"\n🏗️  CONTENT STRUCTURE:")
    print(f"   📝 Introduction: ✅")
    print(f"   📄 Main Sections: {len(main_sections)}")
    
    for i, section in enumerate(main_sections, 1):
        heading = section.get("heading", "N/A")
        subsections = section.get("subsections", [])
        print(f"      {i}. {heading}")
        if subsections:
            for j, sub in enumerate(subsections, 1):
                subheading = sub.get("subheading", "N/A")
                print(f"         {i}.{j} {subheading}")
    
    print(f"   🎯 Conclusion: ✅")
    
    # SEO elements
    seo_elements = blog_post.get("seo_elements", {})
    if seo_elements:
        print(f"\n🔍 SEO OPTIMIZATION:")
        primary_kw = seo_elements.get("primary_keywords_used", [])
        secondary_kw = seo_elements.get("secondary_keywords_used", [])
        
        if primary_kw:
            print(f"   🎯 Primary Keywords Used: {', '.join(primary_kw[:5])}")
        if secondary_kw:
            print(f"   🔑 Secondary Keywords Used: {', '.join(secondary_kw[:5])}")
        
        internal_links = seo_elements.get("internal_link_opportunities", [])
        external_links = seo_elements.get("external_link_opportunities", [])
        
        if internal_links:
            print(f"   🔗 Internal Link Opportunities: {len(internal_links)}")
        if external_links:
            print(f"   🌐 External Link Suggestions: {len(external_links)}")
    
    print(f"\n✅ Blog post generation complete!")


def save_as_docx(formatted_post, filename):
    """
    Save the formatted blog post as a .docx file.
    Handles basic markdown like headings, bold, and italics.
    """
    doc = Document()
    
    for line in formatted_post.split('\n'):
        stripped_line = line.strip()
        if not stripped_line or stripped_line.startswith('<!--'):
            continue

        if stripped_line.startswith('# '):
            doc.add_heading(stripped_line[2:].strip(), level=1)
        elif stripped_line.startswith('## '):
            doc.add_heading(stripped_line[3:].strip(), level=2)
        elif stripped_line.startswith('### '):
            doc.add_heading(stripped_line[4:].strip(), level=3)
        else:
            p = doc.add_paragraph()
            # Split by bold/italic markers to handle inline formatting
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                elif part:
                    p.add_run(part)
    
    doc.save(filename)


def create_complete_blog_workflow(target_keyword, num_competitors=3):
    """
    Complete workflow: Analysis → Outline → Blog Post Generation
    
    Args:
        target_keyword (str): Target keyword to analyze
        num_competitors (int): Number of competitors to analyze
        
    Returns:
        dict: Complete workflow results including blog post
    """
    print(f"\n🚀 COMPLETE BLOG CREATION WORKFLOW")
    print(f"🎯 Target Keyword: '{target_keyword}'")
    print("="*80)
    
    # Step 1: Get complete content strategy
    print(f"\n📊 STEP 1: CREATING CONTENT STRATEGY")
    complete_strategy = create_complete_content_strategy(target_keyword, num_competitors)
    
    if "error" in complete_strategy:
        return complete_strategy
    
    # Step 2: Generate blog post from outline
    print(f"\n📝 STEP 2: GENERATING BLOG POST")
    blog_outline = complete_strategy.get("blog_outline", {})
    analysis_results = complete_strategy.get("competitor_analysis", {})
    
    blog_post = generate_blog_post(blog_outline, analysis_results)
    
    if "error" in blog_post:
        return {"error": f"Blog post generation failed: {blog_post['error']}"}
    
    # Step 3: Format for publication
    print(f"\n📋 STEP 3: FORMATTING FOR PUBLICATION")
    formatted_post = format_blog_post_for_publication(blog_post)
    
    # Combine all results
    complete_workflow = {
        "target_keyword": target_keyword,
        "creation_date": datetime.now().isoformat(),
        "content_strategy": complete_strategy,
        "blog_post": blog_post,
        "formatted_post": formatted_post,
        "workflow_summary": {
            "competitors_analyzed": complete_strategy.get("strategy_summary", {}).get("competitors_analyzed", 0),
            "outline_generated": complete_strategy.get("strategy_summary", {}).get("outline_generated", False),
            "blog_post_created": "error" not in blog_post,
            "estimated_word_count": blog_post.get("blog_post", {}).get("word_count", "N/A")
        }
    }
    
    return complete_workflow




if __name__ == "__main__":
    target_keyword = "how to use ai tools for email marketing"
    
    # Check if required API keys are set
    if not SERP_API_KEY:
        print("❌ Please set SERP_API_KEY in your .env file")
        exit(1)
    
    if not GOOGLE_API_KEY:
        print("❌ Please set GOOGLE_API_KEY in your .env file for Gemini analysis")
        exit(1)
    
    # Choose workflow option
    print("🎯 CONTENT CREATION WORKFLOW OPTIONS:")
    print("1. Complete Blog Workflow (Analysis + Outline + Blog Post)")
    print("2. Generate Blog Post from Existing Outline")
    print("3. Complete Content Strategy (Analysis + Outline)")
    print("4. Generate Blog Outline from Existing Analysis")
    print("5. Competitor Analysis Only")
    
    choice = input("\nEnter your choice (1-5): ").strip()
    
    if choice == "1":
        # Option 1: Complete Blog Workflow (Analysis + Outline + Blog Post)
        print(f"\n🚀 Creating Complete Blog (Analysis → Outline → Blog Post)...")
        complete_workflow = create_complete_blog_workflow(target_keyword, num_competitors=3)
        
        if "error" not in complete_workflow:
            # Display all results
            display_analysis_results(complete_workflow["content_strategy"]["competitor_analysis"])
            display_blog_outline(complete_workflow["content_strategy"]["blog_outline"])
            display_blog_post_summary(complete_workflow["blog_post"])
            
            # Save everything
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            # Save complete workflow
            workflow_filename = f"complete_blog_workflow_{timestamp}.json"
            with open(workflow_filename, 'w', encoding='utf-8') as f:
                json.dump(complete_workflow, f, indent=2, ensure_ascii=False)
            
            # Save formatted blog post as .docx
            blog_filename_docx = f"blog_post_{timestamp}.docx"
            save_as_docx(complete_workflow["formatted_post"], blog_filename_docx)
            
            print(f"\n💾 Complete workflow saved to: {workflow_filename}")
            print(f"📝 Blog post saved to: {blog_filename_docx}")
            print(f"📊 Blog is ready for publication!")
        else:
            print(f"❌ Error: {complete_workflow['error']}")
    
    elif choice == "2":
        # Option 2: Generate Blog Post from Existing Outline
        outline_file = input("Enter the blog outline JSON filename: ").strip()
        analysis_file = input("Enter the analysis JSON filename (optional, press Enter to skip): ").strip()
        
        try:
            # Load outline
            with open(outline_file, 'r', encoding='utf-8') as f:
                blog_outline = json.load(f)
            
            # Load analysis (optional)
            analysis_results = None
            if analysis_file:
                try:
                    with open(analysis_file, 'r', encoding='utf-8') as f:
                        analysis_results = json.load(f)
                except FileNotFoundError:
                    print("⚠️  Analysis file not found, proceeding without it...")
            
            print(f"\n📝 Generating Blog Post from outline...")
            blog_post = generate_blog_post(blog_outline, analysis_results)
            
            if "error" not in blog_post:
                display_blog_post_summary(blog_post)
                
                # Format and save as .docx
                formatted_post = format_blog_post_for_publication(blog_post)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                
                blog_filename_docx = f"blog_post_{timestamp}.docx"
                save_as_docx(formatted_post, blog_filename_docx)
                
                # Save blog post data
                blog_json_filename = f"blog_post_data_{timestamp}.json"
                with open(blog_json_filename, 'w', encoding='utf-8') as f:
                    json.dump(blog_post, f, indent=2, ensure_ascii=False)
                
                print(f"\n📝 Blog post saved to: {blog_filename_docx}")
                print(f"💾 Blog post data saved to: {blog_json_filename}")
            else:
                print(f"❌ Error generating blog post: {blog_post['error']}")
                
        except FileNotFoundError:
            print(f"❌ Outline file not found: {outline_file}")
        except json.JSONDecodeError:
            print(f"❌ Invalid JSON file: {outline_file}")
    
    elif choice == "3":
        # Option 3: Complete Content Strategy (Original workflow)
        print(f"\n🚀 Creating Complete Content Strategy...")
        complete_strategy = create_complete_content_strategy(target_keyword, num_competitors=3)
        
        if "error" not in complete_strategy:
            # Display results
            display_analysis_results(complete_strategy["competitor_analysis"])
            display_blog_outline(complete_strategy["blog_outline"])
            
            # Save complete strategy
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            strategy_filename = f"complete_content_strategy_{timestamp}.json"
            
            with open(strategy_filename, 'w', encoding='utf-8') as f:
                json.dump(complete_strategy, f, indent=2, ensure_ascii=False)
            
            print(f"\n💾 Complete content strategy saved to: {strategy_filename}")
            
            # Ask if user wants to generate blog post
            generate_post = input("\nWould you like to generate a blog post from this strategy? (y/n): ").strip().lower()
            if generate_post == 'y':
                print(f"\n📝 Generating Blog Post...")
                blog_post = generate_blog_post(complete_strategy["blog_outline"], complete_strategy["competitor_analysis"])
                
                if "error" not in blog_post:
                    display_blog_post_summary(blog_post)
                    
                    # Format and save as .docx
                    formatted_post = format_blog_post_for_publication(blog_post)
                    blog_filename_docx = f"blog_post_{timestamp}.docx"
                    save_as_docx(formatted_post, blog_filename_docx)
                    
                    print(f"\n📝 Blog post saved to: {blog_filename_docx}")
                else:
                    print(f"❌ Error generating blog post: {blog_post['error']}")
        else:
            print(f"❌ Error: {complete_strategy['error']}")
    
    elif choice == "4":
        # Option 4: Generate outline from existing analysis (unchanged)
        analysis_file = input("Enter the analysis JSON filename (e.g., competitor_analysis_20250904_155636.json): ").strip()
        
        try:
            with open(analysis_file, 'r', encoding='utf-8') as f:
                analysis_results = json.load(f)
            
            print(f"\n📝 Generating Blog Outline from existing analysis...")
            blog_outline = generate_blog_outline(analysis_results)
            display_blog_outline(blog_outline)
            
            # Save outline
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            outline_filename = f"blog_outline_{timestamp}.json"
            
            with open(outline_filename, 'w', encoding='utf-8') as f:
                json.dump(blog_outline, f, indent=2, ensure_ascii=False)
            
            print(f"\n💾 Blog outline saved to: {outline_filename}")
            
            # Ask if user wants to generate blog post
            generate_post = input("\nWould you like to generate a blog post from this outline? (y/n): ").strip().lower()
            if generate_post == 'y':
                print(f"\n📝 Generating Blog Post...")
                blog_post = generate_blog_post(blog_outline, analysis_results)
                
                if "error" not in blog_post:
                    display_blog_post_summary(blog_post)
                    
                    # Format and save as .docx
                    formatted_post = format_blog_post_for_publication(blog_post)
                    blog_filename_docx = f"blog_post_{timestamp}.docx"
                    save_as_docx(formatted_post, blog_filename_docx)
                    
                    print(f"\n📝 Blog post saved to: {blog_filename_docx}")
                else:
                    print(f"❌ Error generating blog post: {blog_post['error']}")
            
        except FileNotFoundError:
            print(f"❌ File not found: {analysis_file}")
        except json.JSONDecodeError:
            print(f"❌ Invalid JSON file: {analysis_file}")
    
    elif choice == "5":
        # Option 5: Competitor Analysis Only (Original workflow)
        print("🚀 Starting AI-Powered Comprehensive Competitor Analysis...")
        analysis_results = comprehensive_competitor_analysis(target_keyword, num_results=3)
        
        # Display summary results
        display_analysis_results(analysis_results)
        
        # Save complete results to JSON file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"competitor_analysis_{timestamp}.json"
        
        with open(output_filename, 'w', encoding='utf-8') as f:
            json.dump(analysis_results, f, indent=2, ensure_ascii=False)
        
        print(f"\n💾 Complete analysis results saved to: {output_filename}")
        
        # Optional: Also save scraped data to text file (legacy)
        print(f"\n📝 Saving scraped data to text file...")
        links, questions = get_top_5_links_and_paa(target_keyword, SERP_API_KEY)
        
        with open("scraped_blog_data.txt", "w", encoding="utf-8") as f:
            f.write(f"Analysis for keyword: {target_keyword}\n")
            f.write(f"Analysis date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            f.write("People Also Ask Questions:\n")
            for question in questions:
                f.write(f" - {question}\n")
            f.write("\n")
            
            f.write("Competitor Blog Data:\n")
            for link in links:
                blog_data = scrape_blog(link)
                if blog_data:
                    f.write(f"\n--- Blog data for {link} ---\n")
                    for key, value in blog_data.items():
                        f.write(f"  {key}: {value}\n")
                    f.write("\n")
        
        print("✅ Blog data also written to scraped_blog_data.txt")
        
        # Ask if user wants to generate outline
        generate_outline = input("\nWould you like to generate a blog outline from this analysis? (y/n): ").strip().lower()
        if generate_outline == 'y':
            print(f"\n📝 Generating Blog Outline...")
            blog_outline = generate_blog_outline(analysis_results)
            display_blog_outline(blog_outline)
            
            # Save outline
            outline_timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            outline_filename = f"blog_outline_{outline_timestamp}.json"
            
            with open(outline_filename, 'w', encoding='utf-8') as f:
                json.dump(blog_outline, f, indent=2, ensure_ascii=False)
            
            print(f"\n💾 Blog outline saved to: {outline_filename}")
            
            # Ask if user wants to generate blog post
            generate_post = input("\nWould you like to generate a blog post from this outline? (y/n): ").strip().lower()
            if generate_post == 'y':
                print(f"\n📝 Generating Blog Post...")
                blog_post = generate_blog_post(blog_outline, analysis_results)
                
                if "error" not in blog_post:
                    display_blog_post_summary(blog_post)
                    
                    # Format and save as .docx
                    formatted_post = format_blog_post_for_publication(blog_post)
                    blog_filename_docx = f"blog_post_{outline_timestamp}.docx"
                    save_as_docx(formatted_post, blog_filename_docx)
                    
                    print(f"📝 Blog post saved to: {blog_filename_docx}")
                else:
                    print(f"❌ Error generating blog post: {blog_post['error']}")
    
    else:
        print("❌ Invalid choice. Please run the script again and choose 1-5.")
    
    print("\n🎉 All operations complete!")

